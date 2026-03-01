from __future__ import annotations

"""Render enhanced fairness manuscript artifacts (figures, tables, and narrative report)."""

import argparse
from dataclasses import dataclass
from datetime import datetime
import json
import logging
import math
from pathlib import Path
import textwrap
from typing import Dict, Iterable, List, Mapping, Sequence

import matplotlib

# Use a non-interactive backend so artifact generation works in headless environments.
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

LOGGER = logging.getLogger(__name__)

DEFAULT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_ARTIFACTS_DIR = DEFAULT_ROOT / "artifacts" / "fairness_enhanced"
FIGURE_EXPORT_DPI = 360
REPORT_PDF_DPI = 320

_PROXY_FALLBACK = "image-derived perilesional skin tone proxy (ITA)"
_BINARY_FALLBACK = "lower-ITA (darker-appearing) vs higher-ITA (lighter-appearing) strata"

_METRIC_LABELS = {
    "cliffs_delta_iou_lower_vs_higher": "Cliff's delta (IoU)",
    "median_iou_diff_lower_minus_higher": "Median IoU difference (Lower - Higher)",
    "mean_iou_diff_lower_minus_higher": "Mean IoU difference (Lower - Higher)",
    "success_risk_difference_lower_minus_higher": "Success risk difference (Lower - Higher)",
    "success_relative_risk_lower_over_higher": "Success relative risk (Lower/Higher)",
    "success_odds_ratio_lower_over_higher": "Success odds ratio (Lower/Higher)",
}

_COVADJ_METRIC_LABELS = {
    "adjusted_risk_low": "Adjusted success probability (Lower ITA)",
    "adjusted_risk_high": "Adjusted success probability (Higher ITA)",
    "adjusted_rd_low_minus_high": "Adjusted risk difference (Lower - Higher)",
    "adjusted_rr_low_over_high": "Adjusted relative risk (Lower/Higher)",
    "adjusted_or_low_over_high": "Adjusted odds ratio (Lower/Higher)",
    "unadjusted_rd_low_minus_high": "Unadjusted risk difference (Lower - Higher)",
    "rd_attenuation_adj_over_unadj": "RD attenuation (Adjusted/Unadjusted)",
}

_MODEL_CURVE_LABELS = {
    "ita_only": "ITA only",
    "ita_plus_covariates": "ITA + covariates",
}

_COVARIATE_LABELS = {
    "lesion_area_frac": "Lesion area fraction",
    "deltaE_lesion_skin": "Lesion-skin color contrast (Delta E)",
    "Lstar_skin_mean": "Perilesional brightness mean (L*)",
    "Lstar_skin_std": "Perilesional brightness variability (L* SD)",
    "sharpness_laplacian_var": "Image sharpness (variance of Laplacian)",
    "hair_frac": "Hair/occlusion fraction",
    "specular_frac": "Specular highlight fraction",
}

_KEY_TERM_DICTIONARY = [
    (
        "Image-derived perilesional skin tone proxy (ITA)",
        "Computed from median L* and b* values in a perilesional ring around the lesion (excluding lesion pixels).",
        "Higher ITA indicates lighter-appearing perilesional skin in the image; lower ITA indicates darker-appearing appearance strata.",
    ),
    (
        "Lower-ITA vs Higher-ITA strata",
        "Binary grouping of ITA values using the configured cutoff (default 28°).",
        "Defines comparison groups for disparity summaries without inferring patient identity labels.",
    ),
    (
        "Median IoU difference (Lower - Higher)",
        "Difference between subgroup median IoU values.",
        "Negative values indicate lower median segmentation overlap in lower-ITA strata.",
    ),
    (
        "Risk difference (RD)",
        "Difference in success rates (IoU threshold endpoint) between lower- and higher-ITA strata.",
        "Negative RD indicates lower success probability in lower-ITA strata.",
    ),
    (
        "Relative risk (RR)",
        "Ratio of success probabilities: lower-ITA success divided by higher-ITA success.",
        "RR < 1 indicates lower relative success for lower-ITA strata.",
    ),
    (
        "Odds ratio (OR)",
        "Ratio of odds of success in lower-ITA vs higher-ITA strata.",
        "OR < 1 indicates lower odds of success for lower-ITA strata.",
    ),
]


@dataclass(frozen=True)
class EnhancedRunSummary:
    run_id: str
    model_name: str
    prompt_variant: str
    audit_mode: str
    input_pairs: int
    processed_rows: int
    primary_rows: int
    lower_n: int
    higher_n: int
    unknown_n: int
    ita_cutoff: float
    dedup_mode: str
    bootstrap_method: str
    bootstrap_resamples: int
    trend_covariates_used: List[str]


@dataclass(frozen=True)
class FigureSpecEnhanced:
    key: str
    stem: str
    title: str
    caption: str
    supplementary: bool = False


@dataclass(frozen=True)
class TableSpecEnhanced:
    key: str
    stem: str
    title: str
    description: str
    supplementary: bool = False


@dataclass(frozen=True)
class ReportSection:
    title: str
    paragraphs: List[str]
    figure_keys: List[str]
    table_keys: List[str]


@dataclass(frozen=True)
class FigureArtifactSet:
    png: Path
    svg: Path
    pdf: Path


@dataclass(frozen=True)
class TableArtifactSet:
    csv: Path
    html: Path
    markdown: Path


@dataclass(frozen=True)
class EnhancedArtifactBundle:
    figures: Dict[str, FigureArtifactSet]
    tables: Dict[str, TableArtifactSet]
    report_pdf: Path
    report_docx: Path
    report_markdown: Path
    report_html: Path


def _load_json(path: Path) -> Dict[str, object]:
    with path.open("r", encoding="utf-8") as handle:
        payload = json.load(handle)
    if not isinstance(payload, dict):
        raise ValueError(f"Expected JSON object in {path}")
    return payload


def _require_path(path: Path, description: str) -> Path:
    if not path.exists():
        raise FileNotFoundError(f"{description} not found: {path}")
    return path


def _fmt_num(value: object, ndigits: int = 4) -> str:
    try:
        val = float(value)
    except Exception:
        return "NA"
    if math.isnan(val):
        return "NA"
    return f"{val:.{ndigits}f}"


def _fmt_ci(lo: object, hi: object, ndigits: int = 4) -> str:
    try:
        low = float(lo)
        high = float(hi)
    except Exception:
        return "NA"
    if math.isnan(low) or math.isnan(high):
        return "NA"
    return f"[{low:.{ndigits}f}, {high:.{ndigits}f}]"


def _fmt_pct(value: object, ndigits: int = 1) -> str:
    try:
        val = float(value)
    except Exception:
        return "NA"
    if math.isnan(val):
        return "NA"
    return f"{100.0 * val:.{ndigits}f}%"


def _df_to_markdown(df: pd.DataFrame) -> str:
    if df.empty:
        return "| No data |\n|---|\n| NA |"
    cols = [str(c) for c in df.columns]
    lines = []
    lines.append("| " + " | ".join(cols) + " |")
    lines.append("|" + "|".join(["---"] * len(cols)) + "|")
    for _, row in df.iterrows():
        vals = [str(row[c]) for c in df.columns]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def _covariate_display_name(name: str) -> str:
    token = str(name)
    return _COVARIATE_LABELS.get(token, token)


def _covariate_display_list(names: Sequence[str]) -> List[str]:
    return [_covariate_display_name(name) for name in names]


def _wrap_plot_label(value: object, width: int = 36) -> str:
    return textwrap.fill(str(value), width=width, break_long_words=False)


def _save_figure_triplet(fig: plt.Figure, stem: str, output_dir: Path) -> FigureArtifactSet:
    output_dir.mkdir(parents=True, exist_ok=True)
    png = output_dir / f"{stem}.png"
    svg = output_dir / f"{stem}.svg"
    pdf = output_dir / f"{stem}.pdf"
    fig.savefig(png, dpi=FIGURE_EXPORT_DPI, facecolor="white")
    fig.savefig(svg, facecolor="white")
    fig.savefig(pdf, facecolor="white")
    return FigureArtifactSet(png=png, svg=svg, pdf=pdf)


def _write_table_artifacts(df: pd.DataFrame, *, spec: TableSpecEnhanced, output_dir: Path) -> TableArtifactSet:
    output_dir.mkdir(parents=True, exist_ok=True)
    csv_path = output_dir / f"{spec.stem}.csv"
    html_path = output_dir / f"{spec.stem}.html"
    md_path = output_dir / f"{spec.stem}.md"

    df.to_csv(csv_path, index=False)

    html = [
        "<!doctype html>",
        "<html><head><meta charset='utf-8'>",
        f"<title>{spec.title}</title>",
        "<style>body{font-family:Segoe UI,Arial,sans-serif;margin:24px;color:#111;}"
        "table{border-collapse:collapse;width:100%;}"
        "th,td{border:1px solid #d9d9d9;padding:8px 10px;text-align:left;}"
        "th{background:#f5f5f5;} .desc{color:#444; margin-bottom:12px;}</style>",
        "</head><body>",
        f"<h1>{spec.title}</h1>",
        f"<p class='desc'>{spec.description}</p>",
        df.to_html(index=False),
        "</body></html>",
    ]
    html_path.write_text("\n".join(html), encoding="utf-8")

    md = [
        f"# {spec.title}",
        "",
        spec.description,
        "",
        _df_to_markdown(df),
        "",
    ]
    md_path.write_text("\n".join(md), encoding="utf-8")
    return TableArtifactSet(csv=csv_path, html=html_path, markdown=md_path)


def _resolve_primary_frame(df: pd.DataFrame) -> pd.DataFrame:
    if "selected_for_primary" in df.columns:
        out = df[df["selected_for_primary"]].copy()
        if not out.empty:
            return out
    if "is_canonical" in df.columns:
        out = df[df["is_canonical"]].copy()
        if not out.empty:
            return out
    return df.copy()


def _load_enhanced_inputs(fairness_enhanced_dir: Path) -> Dict[str, object]:
    base = fairness_enhanced_dir
    analysis_path = _require_path(base / "analysis_frame.parquet", "analysis_frame.parquet")
    effects_table_path = _require_path(base / "endpoint_effects_table.csv", "endpoint_effects_table.csv")
    effects_json_path = _require_path(base / "endpoint_effects.json", "endpoint_effects.json")
    trend_success_path = _require_path(base / "ita_trend_success.csv", "ita_trend_success.csv")
    trend_iou_summary_path = _require_path(base / "ita_trend_iou_summary.json", "ita_trend_iou_summary.json")
    threshold_table_path = _require_path(base / "threshold_sensitivity_table.csv", "threshold_sensitivity_table.csv")
    metadata_path = _require_path(base / "run_metadata.json", "run_metadata.json")

    analysis_df = pd.read_parquet(analysis_path)
    primary_df = _resolve_primary_frame(analysis_df)
    effects_df = pd.read_csv(effects_table_path)
    effects_payload = _load_json(effects_json_path)
    trend_success_df = pd.read_csv(trend_success_path)
    trend_iou_summary = _load_json(trend_iou_summary_path)
    threshold_df = pd.read_csv(threshold_table_path)
    run_metadata = _load_json(metadata_path)

    optional = {
        "dedup_sensitivity": pd.read_csv(base / "dedup_sensitivity.csv") if (base / "dedup_sensitivity.csv").exists() else None,
        "mask_source_sensitivity": pd.read_csv(base / "mask_source_sensitivity.csv") if (base / "mask_source_sensitivity.csv").exists() else None,
        "dependence_sensitivity": _load_json(base / "dependence_sensitivity.json") if (base / "dependence_sensitivity.json").exists() else None,
        "dedup_report": pd.read_csv(base / "dedup_report.csv") if (base / "dedup_report.csv").exists() else None,
        "ita_bins_table": pd.read_csv(base / "ita_bins_table.csv") if (base / "ita_bins_table.csv").exists() else None,
        "runtime_profile": _load_json(base / "runtime_profile.json") if (base / "runtime_profile.json").exists() else None,
        "ita_trend_iou_png": base / "ita_trend_iou_median.png" if (base / "ita_trend_iou_median.png").exists() else None,
        "covariates_qc_dir": base / "covariates_qc_plots" if (base / "covariates_qc_plots").exists() else None,
        "covadj_effects": pd.read_csv(base / "covadj_success_t050_effects.csv")
        if (base / "covadj_success_t050_effects.csv").exists()
        else None,
        "covadj_payload": _load_json(base / "covadj_success_t050_effects.json")
        if (base / "covadj_success_t050_effects.json").exists()
        else None,
        "covadj_model_spec": _load_json(base / "covadj_model_spec.json")
        if (base / "covadj_model_spec.json").exists()
        else None,
        "covadj_components": pd.read_csv(base / "covadj_component_effects.csv")
        if (base / "covadj_component_effects.csv").exists()
        else None,
    }

    dataset_counts = run_metadata.get("dataset_counts", {}) if isinstance(run_metadata, dict) else {}
    runtime_cfg = run_metadata.get("config", {}).get("runtime", {}) if isinstance(run_metadata, dict) else {}
    dedup_cfg = run_metadata.get("config", {}).get("dedup", {}) if isinstance(run_metadata, dict) else {}
    bootstrap_cfg = run_metadata.get("bootstrap", {}) if isinstance(run_metadata, dict) else {}

    lower_n = int((primary_df.get("ita_binary", pd.Series(dtype=str)) == "Lower ITA").sum())
    higher_n = int((primary_df.get("ita_binary", pd.Series(dtype=str)) == "Higher ITA").sum())
    unknown_n = int(len(primary_df) - lower_n - higher_n)

    run_summary = EnhancedRunSummary(
        run_id=str(run_metadata.get("config", {}).get("run_id", "unknown")) if isinstance(run_metadata, dict) else "unknown",
        model_name=str(run_metadata.get("config", {}).get("model_name", "unknown")) if isinstance(run_metadata, dict) else "unknown",
        prompt_variant=str(run_metadata.get("config", {}).get("prompt_variant", "unknown")) if isinstance(run_metadata, dict) else "unknown",
        audit_mode=str(run_metadata.get("config", {}).get("audit_mode", "enhanced")) if isinstance(run_metadata, dict) else "enhanced",
        input_pairs=int(dataset_counts.get("input_pairs", len(analysis_df))),
        processed_rows=int(dataset_counts.get("processed_rows", len(analysis_df))),
        primary_rows=int(dataset_counts.get("primary_rows", len(primary_df))),
        lower_n=lower_n,
        higher_n=higher_n,
        unknown_n=unknown_n,
        ita_cutoff=float(run_metadata.get("ita_cutoff", 28.0)) if isinstance(run_metadata, dict) else 28.0,
        dedup_mode=str(dedup_cfg.get("mode", "unknown")),
        bootstrap_method=str(bootstrap_cfg.get("method", "unknown")),
        bootstrap_resamples=int(bootstrap_cfg.get("n_resamples", 0) or 0),
        trend_covariates_used=[str(v) for v in run_metadata.get("trend_covariates_used", [])] if isinstance(run_metadata, dict) else [],
    )

    return {
        "analysis": analysis_df,
        "primary": primary_df,
        "effects": effects_df,
        "effects_payload": effects_payload,
        "trend_success": trend_success_df,
        "trend_iou_summary": trend_iou_summary,
        "threshold": threshold_df,
        "run_metadata": run_metadata,
        "run_summary": run_summary,
        "optional": optional,
    }


def _figure_specs(proxy_caption: str) -> Dict[str, FigureSpecEnhanced]:
    return {
        "E1": FigureSpecEnhanced(
            key="E1",
            stem="figureE1_cohort_accountability",
            title="Figure E1. Cohort Accountability",
            caption="Cohort accounting from input pairs through deduped analytic cohort, including ITA strata counts.",
        ),
        "E2": FigureSpecEnhanced(
            key="E2",
            stem="figureE2_primary_effects_forest",
            title="Figure E2. Primary Endpoint Effect Sizes",
            caption="Effect-size summary (continuous IoU and binary success) with confidence intervals where available.",
        ),
        "E3": FigureSpecEnhanced(
            key="E3",
            stem="figureE3_trends_success_iou",
            title="Figure E3. Continuous ITA Trend Models",
            caption=f"Continuous performance trends over the {proxy_caption}",
        ),
        "E4": FigureSpecEnhanced(
            key="E4",
            stem="figureE4_threshold_sensitivity",
            title="Figure E4. Threshold Sensitivity",
            caption="Risk/ratio endpoints across IoU success thresholds.",
        ),
        "E5": FigureSpecEnhanced(
            key="E5",
            stem="figureE5_ita_distribution",
            title="Figure E5. ITA Distribution and Cutoff",
            caption="Distribution of ITA values with binary cutoff annotation.",
        ),
        "ES1": FigureSpecEnhanced(
            key="ES1",
            stem="figureES1_dedup_sensitivity",
            title="Figure ES1. Deduplication Sensitivity",
            caption="Effect-size stability across deduplication modes.",
            supplementary=True,
        ),
        "ES2": FigureSpecEnhanced(
            key="ES2",
            stem="figureES2_mask_source_sensitivity",
            title="Figure ES2. Mask Provenance Sensitivity",
            caption="Effect sizes stratified by mask-source provenance.",
            supplementary=True,
        ),
        "ES3": FigureSpecEnhanced(
            key="ES3",
            stem="figureES3_covariate_qc_overview",
            title="Figure ES3. Covariate QC Overview",
            caption="Overview of covariate distributions used for adjusted analyses.",
            supplementary=True,
        ),
    }


def _table_specs() -> Dict[str, TableSpecEnhanced]:
    return {
        "E1": TableSpecEnhanced(
            key="E1",
            stem="tableE1_cohort_and_provenance",
            title="Table E1. Cohort and Provenance",
            description="Counts pre/post deduplication with source/split/provenance summaries.",
        ),
        "E2": TableSpecEnhanced(
            key="E2",
            stem="tableE2_primary_endpoint_effects",
            title="Table E2. Primary Endpoint Effect Sizes",
            description="Primary fairness endpoints with estimates, confidence intervals, and test p-values.",
        ),
        "E3": TableSpecEnhanced(
            key="E3",
            stem="tableE3_group_descriptives",
            title="Table E3. Group Descriptives",
            description="Descriptive performance summaries by lower- vs higher-ITA strata.",
        ),
        "E4": TableSpecEnhanced(
            key="E4",
            stem="tableE4_trend_model_spec",
            title="Table E4. Trend Model Specification",
            description="Trend model settings, covariates used, and reproducibility settings.",
        ),
        "E5": TableSpecEnhanced(
            key="E5",
            stem="tableE5_covariate_adjusted_success_effects",
            title="Table E5. Covariate-Adjusted Success Effects",
            description=(
                "Predictive-margin adjusted risks and derived effect sizes for success "
                "(IoU threshold endpoint), with bootstrap confidence intervals."
            ),
        ),
        "E6": TableSpecEnhanced(
            key="E6",
            stem="tableE6_adjusted_model_components",
            title="Table E6. Adjusted Model Component Contributions",
            description=(
                "Adjusted logistic-model term contributions shown as odds ratios with 95% "
                "bootstrap confidence intervals and CI-based significance indicators."
            ),
        ),
        "ES1": TableSpecEnhanced(
            key="ES1",
            stem="tableES1_threshold_sensitivity",
            title="Table ES1. Threshold Sensitivity",
            description="Success endpoint sensitivity across IoU thresholds.",
            supplementary=True,
        ),
        "ES2": TableSpecEnhanced(
            key="ES2",
            stem="tableES2_dedup_sensitivity",
            title="Table ES2. Deduplication Sensitivity",
            description="Effect-size shifts under none/exact/near dedup modes.",
            supplementary=True,
        ),
        "ES3": TableSpecEnhanced(
            key="ES3",
            stem="tableES3_mask_source_sensitivity",
            title="Table ES3. Mask Source Sensitivity",
            description="Stratified effects by mask provenance where available.",
            supplementary=True,
        ),
        "ES4": TableSpecEnhanced(
            key="ES4",
            stem="tableES4_runtime_profile",
            title="Table ES4. Runtime and Checkpoint Profile",
            description="Runtime stage timings and checkpoint/restart metadata.",
            supplementary=True,
        ),
    }


def _render_cohort_accountability(run_summary: EnhancedRunSummary) -> plt.Figure:
    labels = [
        "Input pairs",
        "Processed rows",
        "Primary rows",
        "Lower ITA",
        "Higher ITA",
        "Unknown ITA",
    ]
    values = [
        run_summary.input_pairs,
        run_summary.processed_rows,
        run_summary.primary_rows,
        run_summary.lower_n,
        run_summary.higher_n,
        run_summary.unknown_n,
    ]

    fig, ax = plt.subplots(figsize=(9.5, 5.0))
    y = np.arange(len(labels))
    bars = ax.barh(y, values, color=["#355C7D", "#4D7EA8", "#5D9C59", "#7AA95C", "#C27B48", "#888888"])
    ax.set_yticks(y, labels)
    ax.invert_yaxis()
    ax.set_xlabel("Count")
    ax.set_title("Figure E1. Cohort Accountability")
    ax.grid(axis="x", linestyle=":", alpha=0.3)
    for rect, value in zip(bars, values):
        ax.text(rect.get_width() + max(values) * 0.01, rect.get_y() + rect.get_height() / 2.0, str(int(value)), va="center")
    fig.tight_layout()
    return fig


def _render_primary_effects_forest(effects_df: pd.DataFrame) -> plt.Figure:
    work = effects_df.copy()
    work["label"] = work["metric"].map(_METRIC_LABELS).fillna(work["metric"])
    work["label_wrapped"] = work["label"].apply(_wrap_plot_label)

    additive_metrics = [
        "median_iou_diff_lower_minus_higher",
        "mean_iou_diff_lower_minus_higher",
        "success_risk_difference_lower_minus_higher",
        "cliffs_delta_iou_lower_vs_higher",
    ]
    ratio_metrics = [
        "success_relative_risk_lower_over_higher",
        "success_odds_ratio_lower_over_higher",
    ]

    additive = work[work["metric"].isin(additive_metrics)].copy()
    ratio = work[work["metric"].isin(ratio_metrics)].copy()

    fig, axes = plt.subplots(1, 2, figsize=(14.6, 5.6), gridspec_kw={"width_ratios": [1.25, 1.0]})

    if not additive.empty:
        y = np.arange(len(additive))
        est = additive["estimate"].astype(float).to_numpy()
        lo = pd.to_numeric(additive["ci_lower"], errors="coerce").to_numpy(dtype=float)
        hi = pd.to_numeric(additive["ci_upper"], errors="coerce").to_numpy(dtype=float)

        axes[0].axvline(0.0, color="#555555", linestyle="--", linewidth=1)
        for idx in range(len(additive)):
            if math.isnan(lo[idx]) or math.isnan(hi[idx]):
                axes[0].scatter(est[idx], idx, color="#1f77b4", s=40)
            else:
                axes[0].errorbar(
                    est[idx],
                    idx,
                    xerr=np.array([[max(0.0, est[idx] - lo[idx])], [max(0.0, hi[idx] - est[idx])]]),
                    fmt="o",
                    color="#1f77b4",
                    ecolor="#1f77b4",
                    capsize=3,
                )
        axes[0].set_yticks(y, additive["label_wrapped"].tolist())
        axes[0].set_title("Additive effects")
        axes[0].set_xlabel("Estimate (95% CI)")
        axes[0].grid(axis="x", linestyle=":", alpha=0.3)
        axes[0].tick_params(axis="y", labelsize=9)
    else:
        axes[0].text(0.5, 0.5, "No additive effects available", ha="center", va="center")
        axes[0].axis("off")

    if not ratio.empty:
        y2 = np.arange(len(ratio))
        est2 = ratio["estimate"].astype(float).to_numpy()
        lo2 = pd.to_numeric(ratio["ci_lower"], errors="coerce").to_numpy(dtype=float)
        hi2 = pd.to_numeric(ratio["ci_upper"], errors="coerce").to_numpy(dtype=float)

        axes[1].axvline(1.0, color="#555555", linestyle="--", linewidth=1)
        for idx in range(len(ratio)):
            if math.isnan(lo2[idx]) or math.isnan(hi2[idx]):
                axes[1].scatter(est2[idx], idx, color="#9c2f2f", s=40)
            else:
                axes[1].errorbar(
                    est2[idx],
                    idx,
                    xerr=np.array([[max(0.0, est2[idx] - lo2[idx])], [max(0.0, hi2[idx] - est2[idx])]]),
                    fmt="o",
                    color="#9c2f2f",
                    ecolor="#9c2f2f",
                    capsize=3,
                )
        axes[1].set_xscale("log")
        axes[1].set_yticks(y2, ratio["label_wrapped"].tolist())
        axes[1].set_title("Ratio effects")
        axes[1].set_xlabel("Estimate (95% CI, log scale)")
        axes[1].grid(axis="x", linestyle=":", alpha=0.3)
        axes[1].tick_params(axis="y", labelsize=9)
    else:
        axes[1].text(0.5, 0.5, "No ratio effects available", ha="center", va="center")
        axes[1].axis("off")

    fig.suptitle("Figure E2. Primary Endpoint Effect Sizes", y=0.985)
    fig.subplots_adjust(left=0.30, right=0.98, top=0.88, bottom=0.12, wspace=0.95)
    return fig


def _render_trends_success_iou(
    trend_success_df: pd.DataFrame,
    iou_summary: Mapping[str, object],
    *,
    iou_png_path: Path | None,
    covariates_used: Sequence[str],
) -> plt.Figure:
    fig, axes = plt.subplots(1, 2, figsize=(13.5, 5.2))

    if not trend_success_df.empty:
        for model_name, part in trend_success_df.groupby("model"):
            label = _MODEL_CURVE_LABELS.get(str(model_name), str(model_name))
            axes[0].plot(part["ita_deg"], part["pred"], label=label)
            if {"ci_lower", "ci_upper"}.issubset(part.columns):
                low = pd.to_numeric(part["ci_lower"], errors="coerce")
                high = pd.to_numeric(part["ci_upper"], errors="coerce")
                axes[0].fill_between(part["ita_deg"], low, high, alpha=0.16)
        axes[0].set_ylim(0.0, 1.0)
        axes[0].set_xlabel("ITA (degrees)")
        axes[0].set_ylabel("Predicted success probability")
        axes[0].set_title("Predicted success probability vs ITA")
        axes[0].grid(alpha=0.25, linestyle=":")
        axes[0].legend(loc="best")
    else:
        axes[0].text(0.5, 0.5, "No success trend data", ha="center", va="center")
        axes[0].axis("off")

    if iou_png_path is not None and iou_png_path.exists():
        img = plt.imread(iou_png_path)
        axes[1].imshow(img)
        axes[1].axis("off")
        axes[1].set_title("Predicted median IoU vs ITA")
    else:
        axes[1].axis("off")
        axes[1].set_title("Predicted median IoU trend summary")

    base_status = str(((iou_summary.get("base") or {}) if isinstance(iou_summary, dict) else {}).get("status", "unknown"))
    adj_status = str(((iou_summary.get("covariate_adjusted") or {}) if isinstance(iou_summary, dict) else {}).get("status", "unknown"))
    cov_text = "; ".join(_covariate_display_list(covariates_used)) if covariates_used else "none"
    summary = textwrap.fill(
        f"IoU trend summary: base={base_status}; adjusted={adj_status}; covariates={cov_text}",
        width=130,
        break_long_words=False,
    )
    fig.text(
        0.02,
        0.01,
        summary,
        fontsize=9,
        ha="left",
        va="bottom",
    )
    fig.suptitle("Figure E3. Continuous ITA Trend Models", y=0.985)
    fig.tight_layout(rect=(0.02, 0.15, 0.98, 0.95))
    return fig


def _render_threshold_sensitivity(threshold_df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 3, figsize=(14.5, 4.8), sharex=True)

    for ax, metric, color, null_line, panel_title in (
        (axes[0], "rd", "#9c2f2f", 0.0, "Risk difference (RD)"),
        (axes[1], "rr", "#1f77b4", 1.0, "Relative risk (RR)"),
        (axes[2], "or", "#2c8a57", 1.0, "Odds ratio (OR)"),
    ):
        if metric not in threshold_df.columns:
            ax.text(0.5, 0.5, f"No {metric.upper()} data", ha="center", va="center")
            ax.axis("off")
            continue
        x = threshold_df["threshold"].astype(float)
        y = threshold_df[metric].astype(float)
        lo_col = f"{metric}_ci_lower"
        hi_col = f"{metric}_ci_upper"
        ax.plot(x, y, marker="o", color=color)
        if lo_col in threshold_df.columns and hi_col in threshold_df.columns:
            lo = pd.to_numeric(threshold_df[lo_col], errors="coerce").astype(float)
            hi = pd.to_numeric(threshold_df[hi_col], errors="coerce").astype(float)
            ax.fill_between(x, lo, hi, color=color, alpha=0.18)
        ax.axhline(null_line, color="#555555", linestyle="--", linewidth=1)
        ax.grid(alpha=0.25, linestyle=":")
        ax.set_title(panel_title)
        ax.set_xlabel("IoU threshold")

    axes[0].set_ylabel("Estimate")
    fig.suptitle("Figure E4. Threshold Sensitivity", y=0.985)
    fig.tight_layout(rect=(0.02, 0.06, 0.98, 0.95))
    return fig


def _render_ita_distribution(primary_df: pd.DataFrame, cutoff: float) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(9.5, 4.8))

    values = pd.to_numeric(primary_df.get("ita_deg", pd.Series(dtype=float)), errors="coerce")
    values = values[np.isfinite(values)]
    if values.empty:
        ax.text(0.5, 0.5, "No ITA values available", ha="center", va="center")
        ax.axis("off")
        return fig

    lower = pd.to_numeric(primary_df[primary_df.get("ita_binary") == "Lower ITA"].get("ita_deg", pd.Series(dtype=float)), errors="coerce")
    higher = pd.to_numeric(primary_df[primary_df.get("ita_binary") == "Higher ITA"].get("ita_deg", pd.Series(dtype=float)), errors="coerce")

    ax.hist(lower[np.isfinite(lower)], bins=35, alpha=0.7, color="#c27b48", label="Lower ITA")
    ax.hist(higher[np.isfinite(higher)], bins=35, alpha=0.6, color="#4d7ea8", label="Higher ITA")
    ax.axvline(cutoff, color="#bb2222", linestyle="--", linewidth=1.2, label=f"Cutoff {cutoff:.1f}")
    ax.set_xlabel("ITA (degrees)")
    ax.set_ylabel("Count")
    ax.set_title("Figure E5. ITA Distribution and Binary Cutoff")
    ax.grid(axis="y", linestyle=":", alpha=0.25)
    ax.legend(loc="best")
    fig.tight_layout(rect=(0.02, 0.06, 0.98, 0.96))
    return fig


def _render_dedup_sensitivity(dedup_df: pd.DataFrame) -> plt.Figure:
    fig, ax = plt.subplots(figsize=(9.5, 4.8))
    if dedup_df.empty:
        ax.text(0.5, 0.5, "No dedup sensitivity data", ha="center", va="center")
        ax.axis("off")
        return fig

    cols = [c for c in dedup_df.columns if c.endswith("__estimate")]
    if not cols:
        ax.text(0.5, 0.5, "No estimate columns in dedup sensitivity", ha="center", va="center")
        ax.axis("off")
        return fig

    melted = dedup_df.melt(id_vars=["dedup_mode"], value_vars=cols, var_name="metric", value_name="estimate")
    for metric_name, part in melted.groupby("metric"):
        token = metric_name.replace("__estimate", "")
        label = _METRIC_LABELS.get(token, token)
        ax.plot(part["dedup_mode"], part["estimate"], marker="o", label=label)
    ax.axhline(0.0, color="#555555", linestyle="--", linewidth=1)
    ax.set_title("Figure ES1. Deduplication Sensitivity")
    ax.set_ylabel("Estimate")
    ax.grid(axis="y", linestyle=":", alpha=0.25)
    ax.legend(loc="best", fontsize=8)
    fig.tight_layout(rect=(0.02, 0.06, 0.98, 0.96))
    return fig


def _render_mask_source_sensitivity(mask_df: pd.DataFrame) -> plt.Figure:
    fig, axes = plt.subplots(1, 2, figsize=(13.0, 4.8))
    if mask_df.empty:
        for ax in axes:
            ax.text(0.5, 0.5, "No mask-source sensitivity data", ha="center", va="center")
            ax.axis("off")
        return fig

    x = np.arange(len(mask_df))
    labels = mask_df["mask_source"].astype(str).tolist()

    axes[0].bar(x, pd.to_numeric(mask_df.get("median_iou_diff", pd.Series(dtype=float)), errors="coerce"), color="#4d7ea8")
    axes[0].axhline(0.0, color="#555555", linestyle="--", linewidth=1)
    axes[0].set_xticks(x, labels, rotation=20, ha="right")
    axes[0].set_title("Median IoU difference")
    axes[0].set_ylabel("Estimate")
    axes[0].grid(axis="y", linestyle=":", alpha=0.25)

    axes[1].bar(x, pd.to_numeric(mask_df.get("rd", pd.Series(dtype=float)), errors="coerce"), color="#c27b48")
    axes[1].axhline(0.0, color="#555555", linestyle="--", linewidth=1)
    axes[1].set_xticks(x, labels, rotation=20, ha="right")
    axes[1].set_title("Risk difference")
    axes[1].grid(axis="y", linestyle=":", alpha=0.25)

    fig.suptitle("Figure ES2. Mask Source Sensitivity", y=0.985)
    fig.tight_layout(rect=(0.02, 0.06, 0.98, 0.95))
    return fig


def _render_covariate_qc_overview(qc_dir: Path | None) -> plt.Figure | None:
    if qc_dir is None or not qc_dir.exists():
        return None
    images = sorted(qc_dir.glob("*_hist.png"))
    if not images:
        return None

    images = images[:6]
    n_cols = 3
    n_rows = int(math.ceil(len(images) / float(n_cols)))
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(13.5, 3.8 * n_rows))
    axes_flat = np.array(axes).reshape(-1)

    for ax in axes_flat:
        ax.axis("off")

    for ax, path in zip(axes_flat, images):
        ax.imshow(plt.imread(path))
        token = path.stem
        if token.endswith("_hist"):
            token = token[: -len("_hist")]
        ax.set_title(_covariate_display_name(token), fontsize=10)
        ax.axis("off")

    fig.suptitle("Figure ES3. Covariate QC Overview", y=0.985)
    fig.tight_layout(rect=(0.02, 0.05, 0.98, 0.95))
    return fig


def _render_covariate_qc_overview_from_frame(primary_df: pd.DataFrame) -> plt.Figure | None:
    covariate_order = [
        "deltaE_lesion_skin",
        "hair_frac",
        "lesion_area_frac",
        "Lstar_skin_mean",
        "Lstar_skin_std",
        "sharpness_laplacian_var",
        "specular_frac",
    ]
    available = []
    for col in covariate_order:
        if col not in primary_df.columns:
            continue
        values = pd.to_numeric(primary_df[col], errors="coerce")
        values = values[np.isfinite(values)]
        if values.empty:
            continue
        available.append((col, values))

    if not available:
        return None

    n_show = min(6, len(available))
    n_cols = 3
    n_rows = int(math.ceil(n_show / float(n_cols)))
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(13.5, 3.7 * n_rows))
    axes_flat = np.array(axes).reshape(-1)

    for ax in axes_flat:
        ax.axis("off")

    for ax, (col, values) in zip(axes_flat, available[:n_show]):
        ax.axis("on")
        ax.hist(values.to_numpy(dtype=float), bins=36, color="#5f7f9f", edgecolor="white")
        label = _covariate_display_name(col)
        ax.set_title(label, fontsize=10)
        ax.set_xlabel(label, fontsize=8)
        ax.set_ylabel("Count", fontsize=8)
        ax.tick_params(axis="both", labelsize=7)
        ax.grid(axis="y", linestyle=":", alpha=0.22)

    fig.suptitle("Figure ES3. Covariate QC Overview", y=0.985)
    fig.tight_layout(rect=(0.02, 0.05, 0.98, 0.95))
    return fig


def _table_cohort_and_provenance(primary_df: pd.DataFrame, run_summary: EnhancedRunSummary, dedup_report: pd.DataFrame | None) -> pd.DataFrame:
    rows: List[Dict[str, object]] = [
        {"section": "cohort", "metric": "input_pairs", "value": run_summary.input_pairs},
        {"section": "cohort", "metric": "processed_rows", "value": run_summary.processed_rows},
        {"section": "cohort", "metric": "primary_rows", "value": run_summary.primary_rows},
        {"section": "cohort", "metric": "lower_ita_n", "value": run_summary.lower_n},
        {"section": "cohort", "metric": "higher_ita_n", "value": run_summary.higher_n},
        {"section": "cohort", "metric": "unknown_ita_n", "value": run_summary.unknown_n},
    ]

    if "dataset_source_primary" in primary_df.columns:
        for source, count in primary_df["dataset_source_primary"].fillna("unknown").value_counts().items():
            rows.append({"section": "dataset_source", "metric": str(source), "value": int(count)})
    if "split" in primary_df.columns:
        for split, count in primary_df["split"].fillna("unknown").value_counts().items():
            rows.append({"section": "split", "metric": str(split), "value": int(count)})
    if "mask_source" in primary_df.columns:
        for source, count in primary_df["mask_source"].fillna("unknown").value_counts().items():
            rows.append({"section": "mask_source", "metric": str(source), "value": int(count)})

    if dedup_report is not None and not dedup_report.empty:
        for _, row in dedup_report.iterrows():
            rows.append(
                {
                    "section": "dedup_report",
                    "metric": str(row.get("dataset_source_primary", "unknown")),
                    "value": f"total={int(row.get('total_images', 0))}, canonical={int(row.get('canonical_images', 0))}, collapsed={int(row.get('collapsed_duplicates', 0))}",
                }
            )

    return pd.DataFrame(rows)


def _table_primary_effects(effects_df: pd.DataFrame) -> pd.DataFrame:
    out = effects_df.copy()
    out["metric_label"] = out["metric"].map(_METRIC_LABELS).fillna(out["metric"])
    out["estimate_fmt"] = out["estimate"].apply(_fmt_num)
    out["ci_fmt"] = [
        _fmt_ci(lo, hi)
        for lo, hi in zip(
            out.get("ci_lower", pd.Series(dtype=float)),
            out.get("ci_upper", pd.Series(dtype=float)),
        )
    ]
    out["p_value_fmt"] = out.get("p_value", pd.Series(dtype=float)).apply(_fmt_num)
    cols = ["metric_label", "estimate_fmt", "ci_fmt", "ci_method", "p_value_fmt"]
    out = out[cols].rename(
        columns={
            "metric_label": "metric",
            "estimate_fmt": "estimate",
            "ci_fmt": "95% ci",
            "ci_method": "ci method",
            "p_value_fmt": "p value",
        }
    )
    return out


def _table_group_descriptives(primary_df: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for group_name in ["Lower ITA", "Higher ITA"]:
        part = primary_df[primary_df.get("ita_binary") == group_name]
        iou = pd.to_numeric(part.get("iou", pd.Series(dtype=float)), errors="coerce")
        if iou.empty:
            rows.append(
                {
                    "group": group_name,
                    "n": 0,
                    "mean_iou": "NA",
                    "median_iou": "NA",
                    "iqr_iou": "NA",
                    "failure_rate": "NA",
                    "success_rate": "NA",
                }
            )
            continue
        iou_clean = iou[np.isfinite(iou)]
        q1 = float(np.percentile(iou_clean, 25)) if not iou_clean.empty else math.nan
        q3 = float(np.percentile(iou_clean, 75)) if not iou_clean.empty else math.nan
        rows.append(
            {
                "group": group_name,
                "n": int(len(part)),
                "mean_iou": _fmt_num(np.nanmean(iou_clean) if not iou_clean.empty else math.nan),
                "median_iou": _fmt_num(np.nanmedian(iou_clean) if not iou_clean.empty else math.nan),
                "iqr_iou": _fmt_ci(q1, q3),
                "failure_rate": _fmt_pct(np.mean(iou_clean == 0.0) if not iou_clean.empty else math.nan),
                "success_rate": _fmt_pct(np.mean(iou_clean >= 0.5) if not iou_clean.empty else math.nan),
            }
        )

    return pd.DataFrame(rows)


def _table_trend_model_spec(run_metadata: Mapping[str, object], run_summary: EnhancedRunSummary) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    config = run_metadata.get("config", {}) if isinstance(run_metadata, dict) else {}
    trend_cfg = config.get("trends", {}) if isinstance(config, dict) else {}
    bootstrap_cfg = run_metadata.get("bootstrap", {}) if isinstance(run_metadata, dict) else {}

    rows.append({"parameter": "runtime_stage", "value": run_metadata.get("runtime_stage", "unknown")})
    rows.append({"parameter": "dedup_mode", "value": run_summary.dedup_mode})
    rows.append({"parameter": "ita_cutoff", "value": _fmt_num(run_summary.ita_cutoff, ndigits=1)})
    rows.append({"parameter": "bootstrap_method", "value": bootstrap_cfg.get("method", "unknown")})
    rows.append({"parameter": "bootstrap_resamples", "value": int(bootstrap_cfg.get("n_resamples", 0) or 0)})
    rows.append({"parameter": "trend_knots", "value": int(trend_cfg.get("knots", 0) or 0)})
    rows.append({"parameter": "trend_degree", "value": int(trend_cfg.get("degree", 0) or 0)})
    rows.append({"parameter": "trend_bootstrap_resamples", "value": int(trend_cfg.get("bootstrap_resamples", 0) or 0)})
    rows.append({"parameter": "trend_quantile", "value": _fmt_num(trend_cfg.get("quantile", math.nan), ndigits=3)})
    rows.append({"parameter": "trend_quantile_alpha", "value": _fmt_num(trend_cfg.get("quantile_alpha", math.nan), ndigits=6)})
    rows.append(
        {
            "parameter": "trend_covariates_used",
            "value": "; ".join(_covariate_display_list(run_summary.trend_covariates_used))
            if run_summary.trend_covariates_used
            else "none",
        }
    )
    return pd.DataFrame(rows)


def _table_covadj_effects(
    covadj_df: pd.DataFrame,
    covadj_payload: Mapping[str, object] | None,
) -> pd.DataFrame:
    out = covadj_df.copy()
    if "metric" not in out.columns:
        out["metric"] = "unknown_metric"
    if "estimate" not in out.columns:
        out["estimate"] = math.nan
    out["metric_label"] = out["metric"].map(_COVADJ_METRIC_LABELS).fillna(out["metric"])
    out["estimate_fmt"] = out["estimate"].apply(_fmt_num)
    out["ci_fmt"] = [
        _fmt_ci(lo, hi)
        for lo, hi in zip(
            out.get("ci_lower", pd.Series(dtype=float)),
            out.get("ci_upper", pd.Series(dtype=float)),
        )
    ]
    out["ci_method_fmt"] = out.get("ci_method", pd.Series(["NA"] * len(out))).fillna("NA").astype(str)
    out["significant_fmt"] = "NA"
    metric_tokens = out["metric"].astype(str).tolist()
    sig_vals: List[str] = []
    for idx, metric in enumerate(metric_tokens):
        lo = pd.to_numeric(out.iloc[idx].get("ci_lower"), errors="coerce")
        hi = pd.to_numeric(out.iloc[idx].get("ci_upper"), errors="coerce")
        if not np.isfinite(lo) or not np.isfinite(hi):
            sig_vals.append("NA")
            continue
        if metric.endswith("_rr_low_over_high") or metric.endswith("_or_low_over_high"):
            sig_vals.append("Yes" if (float(lo) > 1.0 or float(hi) < 1.0) else "No")
        elif metric.endswith("_risk_low") or metric.endswith("_risk_high"):
            sig_vals.append("NA")
        else:
            sig_vals.append("Yes" if (float(lo) > 0.0 or float(hi) < 0.0) else "No")
    out["significant_fmt"] = sig_vals

    cols = [
        "metric_label",
        "estimate_fmt",
        "ci_fmt",
        "significant_fmt",
        "ci_method_fmt",
    ]
    table = out[cols].rename(
        columns={
            "metric_label": "metric",
            "estimate_fmt": "estimate",
            "ci_fmt": "95% ci",
            "significant_fmt": "significant (95% ci)",
            "ci_method_fmt": "ci method",
        }
    )

    summary = covadj_payload.get("summary", {}) if isinstance(covadj_payload, dict) else {}
    if isinstance(summary, dict) and summary:
        footer = {
            "metric": "model_summary",
            "estimate": (
                f"r_low={_fmt_num(summary.get('r_low_adj'))}; "
                f"r_high={_fmt_num(summary.get('r_high_adj'))}; "
                f"rd={_fmt_num(summary.get('rd_adj'))}; "
                f"rd_ci={_fmt_ci(summary.get('rd_ci_lower'), summary.get('rd_ci_upper'))}"
            ),
            "95% ci": "NA",
            "significant (95% ci)": "NA",
            "ci method": "json_summary",
        }
        table = pd.concat([table, pd.DataFrame([footer])], ignore_index=True)

    return table


def _table_covadj_components(component_df: pd.DataFrame) -> pd.DataFrame:
    out = component_df.copy()
    if out.empty:
        return pd.DataFrame(
            columns=[
                "component",
                "odds ratio",
                "95% ci",
                "significant (95% ci)",
                "direction",
                "scale",
                "bootstrap p",
            ]
        )
    out["component_fmt"] = out.get("component", pd.Series(dtype=str)).fillna("unknown").astype(str)
    out["or_fmt"] = out.get("or_estimate", pd.Series(dtype=float)).apply(_fmt_num)
    out["or_ci_fmt"] = [
        _fmt_ci(lo, hi)
        for lo, hi in zip(
            out.get("or_ci_lower", pd.Series(dtype=float)),
            out.get("or_ci_upper", pd.Series(dtype=float)),
        )
    ]
    sig = out.get("or_significant_95ci", pd.Series(dtype=str)).fillna("na").astype(str).str.lower()
    out["sig_fmt"] = sig.map({"yes": "Yes", "no": "No"}).fillna("NA")
    out["direction_fmt"] = out.get("direction", pd.Series(dtype=str)).fillna("uncertain").astype(str)
    out["scale_fmt"] = out.get("scale", pd.Series(dtype=str)).fillna("NA").astype(str)
    out["p_boot_fmt"] = out.get("bootstrap_p_two_sided", pd.Series(dtype=float)).apply(_fmt_num)

    prefer = [
        "Lower ITA indicator (vs Higher ITA)",
    ]
    out["_order"] = out["component_fmt"].apply(lambda x: prefer.index(x) if x in prefer else 10_000)
    out = out.sort_values(["_order", "component_fmt"]).drop(columns=["_order"])
    return out[
        [
            "component_fmt",
            "or_fmt",
            "or_ci_fmt",
            "sig_fmt",
            "direction_fmt",
            "scale_fmt",
            "p_boot_fmt",
        ]
    ].rename(
        columns={
            "component_fmt": "component",
            "or_fmt": "odds ratio",
            "or_ci_fmt": "95% ci",
            "sig_fmt": "significant (95% ci)",
            "direction_fmt": "direction",
            "scale_fmt": "scale",
            "p_boot_fmt": "bootstrap p",
        }
    )


def _table_runtime_profile(run_metadata: Mapping[str, object], runtime_profile: Mapping[str, object] | None) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []

    checkpoint = run_metadata.get("checkpoint", {}) if isinstance(run_metadata, dict) else {}
    runtime = run_metadata.get("runtime", {}) if isinstance(run_metadata, dict) else {}
    for key in [
        "checkpoint_every",
        "processed_this_run",
        "skipped_from_resume",
        "checkpoints_written",
        "manifest_restarts",
        "shard_count",
    ]:
        if key in checkpoint:
            rows.append({"scope": "checkpoint", "parameter": key, "value": checkpoint[key]})

    if isinstance(runtime, dict):
        for key in [
            "requested_workers",
            "workers",
            "workers_auto",
            "max_inflight_tasks_effective",
            "checkpoint_every",
        ]:
            if key in runtime:
                rows.append({"scope": "runtime", "parameter": key, "value": runtime[key]})
        wall = runtime.get("stage_wall_seconds", {})
        if isinstance(wall, dict):
            for stage_name, seconds in wall.items():
                rows.append(
                    {
                        "scope": "stage_wall_seconds",
                        "parameter": str(stage_name),
                        "value": _fmt_num(seconds, ndigits=3),
                    }
                )

    if isinstance(runtime_profile, dict):
        stages = runtime_profile.get("stages", {})
        if isinstance(stages, dict):
            for stage_name, payload in stages.items():
                if not isinstance(payload, dict):
                    continue
                rows.append(
                    {
                        "scope": "runtime_profile",
                        "parameter": f"{stage_name}.throughput_items_per_sec",
                        "value": _fmt_num(payload.get("throughput_items_per_sec", math.nan), ndigits=4),
                    }
                )
                rows.append(
                    {
                        "scope": "runtime_profile",
                        "parameter": f"{stage_name}.rss_peak_bytes",
                        "value": payload.get("rss_peak_bytes", "NA"),
                    }
                )

    return pd.DataFrame(rows)


def _build_sections(
    *,
    report_title: str,
    run_summary: EnhancedRunSummary,
    run_metadata: Mapping[str, object],
    endpoint_payload: Mapping[str, object],
    effects_df: pd.DataFrame,
    covadj_df: pd.DataFrame | None,
    covadj_components_df: pd.DataFrame | None,
    covadj_payload: Mapping[str, object] | None,
    include_supplement: bool,
    available_supp_figures: Sequence[str],
    available_supp_tables: Sequence[str],
) -> List[ReportSection]:
    label_text = endpoint_payload.get("label_text", {}) if isinstance(endpoint_payload, dict) else {}
    methods_snippet = str(label_text.get("methods_snippet", f"Analyses used the {_PROXY_FALLBACK}."))

    med_row = effects_df[effects_df["metric"] == "median_iou_diff_lower_minus_higher"]
    rd_row = effects_df[effects_df["metric"] == "success_risk_difference_lower_minus_higher"]
    med_est = _fmt_num(float(med_row.iloc[0]["estimate"])) if not med_row.empty else "NA"
    rd_est = _fmt_num(float(rd_row.iloc[0]["estimate"])) if not rd_row.empty else "NA"
    covadj_available = isinstance(covadj_df, pd.DataFrame) and not covadj_df.empty
    covadj_components_available = isinstance(covadj_components_df, pd.DataFrame) and not covadj_components_df.empty

    covadj_summary_line = "Covariate-adjusted success model outputs were not available in this run/stage."
    if covadj_available and "metric" in covadj_df.columns:
        rd_adj_row = covadj_df[covadj_df["metric"] == "adjusted_rd_low_minus_high"]
        rr_adj_row = covadj_df[covadj_df["metric"] == "adjusted_rr_low_over_high"]
        if not rd_adj_row.empty:
            rd_val = _fmt_num(rd_adj_row.iloc[0].get("estimate"))
            rd_ci = _fmt_ci(rd_adj_row.iloc[0].get("ci_lower"), rd_adj_row.iloc[0].get("ci_upper"))
            covadj_summary_line = f"Covariate-adjusted success RD (Lower - Higher): {rd_val} ({rd_ci})."
        if not rr_adj_row.empty:
            rr_val = _fmt_num(rr_adj_row.iloc[0].get("estimate"))
            rr_ci = _fmt_ci(rr_adj_row.iloc[0].get("ci_lower"), rr_adj_row.iloc[0].get("ci_upper"))
            covadj_summary_line += f" Adjusted RR: {rr_val} ({rr_ci})."

    covadj_interpret_line = ""
    summary_obj = covadj_payload.get("summary", {}) if isinstance(covadj_payload, dict) else {}
    if isinstance(summary_obj, dict) and summary_obj:
        rd_lo = summary_obj.get("rd_ci_lower")
        rd_hi = summary_obj.get("rd_ci_upper")
        rd_lo_f = pd.to_numeric(rd_lo, errors="coerce")
        rd_hi_f = pd.to_numeric(rd_hi, errors="coerce")
        if np.isfinite(rd_lo_f) and np.isfinite(rd_hi_f):
            if float(rd_lo_f) > 0.0 or float(rd_hi_f) < 0.0:
                covadj_interpret_line = (
                    "Adjusted disparity signal persisted by RD criterion (95% CI excludes 0); "
                    "interpret descriptively under proxy/non-causal framing."
                )
            else:
                covadj_interpret_line = (
                    "Adjusted RD 95% CI included 0, indicating attenuation/uncertainty after covariate adjustment "
                    "under this model specification."
                )
    primary_paragraphs = [
        f"Primary effect estimates showed median IoU difference (lower minus higher) of {med_est} and success risk difference of {rd_est}.",
        "Confidence intervals and supporting nonparametric test p-values are reported alongside each endpoint.",
        covadj_summary_line,
    ]
    if covadj_interpret_line:
        primary_paragraphs.append(covadj_interpret_line)

    sensitivity_text = (
        "Supplementary sensitivity analyses were generated and are included below."
        if include_supplement and (available_supp_figures or available_supp_tables)
        else "Sensitivity supplements were not available for this run stage; core findings should be interpreted with this limitation."
    )
    covariate_labels = _covariate_display_list(run_summary.trend_covariates_used)
    covariate_text = ", ".join(covariate_labels) if covariate_labels else "none"

    dictionary_lines: List[str] = []
    for term, derived, meaning in _KEY_TERM_DICTIONARY:
        dictionary_lines.append(f"{term}: Derived as {derived} Interpretation: {meaning}")
    if covariate_labels:
        dictionary_lines.append(
            "Covariates used in adjusted trends: "
            + "; ".join(covariate_labels)
            + "."
        )
    for raw in run_summary.trend_covariates_used:
        if raw not in _COVARIATE_LABELS:
            continue
        label = _COVARIATE_LABELS[raw]
        if raw == "lesion_area_frac":
            dictionary_lines.append(
                f"{label}: Derived as lesion pixels divided by total image pixels. Interpretation: controls for lesion-size visibility differences."
            )
        elif raw == "deltaE_lesion_skin":
            dictionary_lines.append(
                f"{label}: Derived from Lab color-distance between lesion and perilesional skin summaries. Interpretation: captures lesion-to-skin contrast."
            )
        elif raw == "Lstar_skin_mean":
            dictionary_lines.append(
                f"{label}: Derived as mean perilesional L* brightness. Interpretation: exposure/illumination level proxy."
            )
        elif raw == "Lstar_skin_std":
            dictionary_lines.append(
                f"{label}: Derived as perilesional L* standard deviation. Interpretation: illumination heterogeneity proxy."
            )
        elif raw == "sharpness_laplacian_var":
            dictionary_lines.append(
                f"{label}: Derived from variance of image Laplacian response. Interpretation: focus/blur proxy."
            )
        elif raw == "hair_frac":
            dictionary_lines.append(
                f"{label}: Derived from a hair-like structure detector over lesion/perilesional regions. Interpretation: occlusion burden proxy."
            )
        elif raw == "specular_frac":
            dictionary_lines.append(
                f"{label}: Derived from high-brightness, low-chroma highlight detection. Interpretation: glare burden proxy."
            )

    return [
        ReportSection(
            title="Key Terms and Interpretation Guide",
            paragraphs=dictionary_lines,
            figure_keys=[],
            table_keys=[],
        ),
        ReportSection(
            title="Methods Snapshot",
            paragraphs=[
                methods_snippet,
                "This report is descriptive and non-causal; proxy strata reflect image appearance rather than patient identity attributes.",
            ],
            figure_keys=[],
            table_keys=[],
        ),
        ReportSection(
            title="Cohort and Integrity",
            paragraphs=[
                f"Input pairs: {run_summary.input_pairs}; processed rows: {run_summary.processed_rows}; analytic cohort after dedup selection: {run_summary.primary_rows}.",
                f"Binary strata counts at ITA cutoff {run_summary.ita_cutoff:.1f}: lower={run_summary.lower_n}, higher={run_summary.higher_n}, unknown={run_summary.unknown_n}.",
            ],
            figure_keys=["E1"],
            table_keys=["E1"],
        ),
        ReportSection(
            title="Primary Results",
            paragraphs=primary_paragraphs,
            figure_keys=["E2"],
            table_keys=["E2", "E3"] + (["E5"] if covadj_available else []) + (["E6"] if covadj_components_available else []),
        ),
        ReportSection(
            title="Trend Interpretation",
            paragraphs=[
                "Continuous ITA trend curves are shown for ITA-only and covariate-adjusted formulations where available.",
                f"Covariates used in this run: {covariate_text}.",
                "Threshold sensitivity contextualizes endpoint stability across common IoU decision cutoffs.",
            ],
            figure_keys=["E3", "E4", "E5"],
            table_keys=["E4"],
        ),
        ReportSection(
            title="Sensitivity Results",
            paragraphs=[sensitivity_text],
            figure_keys=[k for k in ["ES1", "ES2", "ES3"] if k in set(available_supp_figures)],
            table_keys=[k for k in ["ES1", "ES2", "ES3", "ES4"] if k in set(available_supp_tables)],
        ),
        ReportSection(
            title="Limitations",
            paragraphs=[
                "Proxy-based fairness analyses should not be interpreted as direct biological or demographic labels.",
                "Sensitivity artifacts may be unavailable in core-only runs and should be generated when full robustness characterization is required.",
            ],
            figure_keys=[],
            table_keys=[],
        ),
        ReportSection(
            title="Reproducibility Appendix",
            paragraphs=[
                f"Report title: {report_title}",
                f"Generated at: {datetime.now().isoformat(timespec='seconds')}",
                f"Runtime stage: {run_metadata.get('runtime_stage', 'unknown') if isinstance(run_metadata, dict) else 'unknown'}",
            ],
            figure_keys=[],
            table_keys=[],
        ),
    ]


def _render_report_markdown(
    *,
    report_title: str,
    sections: Sequence[ReportSection],
    figure_specs: Mapping[str, FigureSpecEnhanced],
    figure_artifacts: Mapping[str, FigureArtifactSet],
    table_specs: Mapping[str, TableSpecEnhanced],
    table_artifacts: Mapping[str, TableArtifactSet],
    output_path: Path,
) -> None:
    lines = [
        f"# {report_title}",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
    ]

    for section in sections:
        lines.append(f"## {section.title}")
        lines.append("")
        for para in section.paragraphs:
            lines.append(para)
            lines.append("")

        for fig_key in section.figure_keys:
            spec = figure_specs[fig_key]
            art = figure_artifacts.get(fig_key)
            if art is None:
                lines.append(f"- {spec.title}: Not available in this run/stage.")
                continue
            rel = Path("figures") / art.png.name
            lines.append(f"### {spec.title}")
            lines.append("")
            lines.append(spec.caption)
            lines.append("")
            lines.append(f"![{spec.title}]({rel.as_posix()})")
            lines.append("")

        for tbl_key in section.table_keys:
            spec = table_specs[tbl_key]
            art = table_artifacts.get(tbl_key)
            if art is None:
                lines.append(f"- {spec.title}: Not available in this run/stage.")
                continue
            lines.append(f"### {spec.title}")
            lines.append("")
            lines.append(spec.description)
            lines.append("")
            lines.append(art.markdown.read_text(encoding="utf-8"))
            lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")


def _render_report_html(
    *,
    report_title: str,
    sections: Sequence[ReportSection],
    figure_specs: Mapping[str, FigureSpecEnhanced],
    figure_artifacts: Mapping[str, FigureArtifactSet],
    table_specs: Mapping[str, TableSpecEnhanced],
    table_artifacts: Mapping[str, TableArtifactSet],
    output_path: Path,
) -> None:
    html: List[str] = [
        "<!doctype html>",
        "<html><head><meta charset='utf-8'>",
        f"<title>{report_title}</title>",
        "<style>body{font-family:Segoe UI,Arial,sans-serif;margin:24px;color:#111;}"
        "h1{margin-bottom:0;} .meta{color:#444;margin-bottom:18px;}"
        "img{max-width:100%;height:auto;border:1px solid #ddd;padding:4px;background:#fff;}"
        "table{border-collapse:collapse;width:100%;margin-top:8px;}"
        "th,td{border:1px solid #d9d9d9;padding:6px 8px;text-align:left;}"
        "th{background:#f5f5f5;} .na{color:#777;font-style:italic;}</style>",
        "</head><body>",
        f"<h1>{report_title}</h1>",
        f"<div class='meta'>Generated: {datetime.now().isoformat(timespec='seconds')}</div>",
    ]

    for section in sections:
        html.append(f"<h2>{section.title}</h2>")
        for para in section.paragraphs:
            html.append(f"<p>{para}</p>")

        for fig_key in section.figure_keys:
            spec = figure_specs[fig_key]
            art = figure_artifacts.get(fig_key)
            if art is None:
                html.append(f"<p class='na'>{spec.title}: Not available in this run/stage.</p>")
                continue
            rel = Path("figures") / art.png.name
            html.append(f"<h3>{spec.title}</h3>")
            html.append(f"<p>{spec.caption}</p>")
            html.append(f"<img src='{rel.as_posix()}' alt='{spec.title}' />")

        for tbl_key in section.table_keys:
            spec = table_specs[tbl_key]
            art = table_artifacts.get(tbl_key)
            if art is None:
                html.append(f"<p class='na'>{spec.title}: Not available in this run/stage.</p>")
                continue
            html.append(f"<h3>{spec.title}</h3>")
            html.append(f"<p>{spec.description}</p>")
            tbl_html = art.html.read_text(encoding="utf-8")
            body_start = tbl_html.find("<table")
            body_end = tbl_html.rfind("</table>")
            if body_start >= 0 and body_end > body_start:
                html.append(tbl_html[body_start : body_end + len("</table>")])
            else:
                html.append("<p class='na'>Unable to embed table HTML.</p>")

    html.append("</body></html>")
    output_path.write_text("\n".join(html), encoding="utf-8")


def _docx_add_table(document: Document, df: pd.DataFrame, *, max_rows: int = 30) -> None:
    shown = df.head(max_rows)
    table = document.add_table(rows=1, cols=len(shown.columns))
    for idx, col in enumerate(shown.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    if len(df) > len(shown):
        document.add_paragraph(f"(Truncated to first {len(shown)} rows for DOCX display)")


def _render_report_docx(
    *,
    report_title: str,
    sections: Sequence[ReportSection],
    figure_specs: Mapping[str, FigureSpecEnhanced],
    figure_artifacts: Mapping[str, FigureArtifactSet],
    table_specs: Mapping[str, TableSpecEnhanced],
    table_artifacts: Mapping[str, TableArtifactSet],
    output_path: Path,
) -> None:
    document = Document()
    document.add_heading(report_title, 0)
    document.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")

    for section in sections:
        document.add_heading(section.title, level=1)
        for para in section.paragraphs:
            document.add_paragraph(para)

        for fig_key in section.figure_keys:
            spec = figure_specs[fig_key]
            art = figure_artifacts.get(fig_key)
            if art is None:
                document.add_paragraph(f"{spec.title}: Not available in this run/stage.")
                continue
            document.add_paragraph(f"{spec.title}. {spec.caption}")
            document.add_picture(str(art.png), width=Inches(6.5))

        for tbl_key in section.table_keys:
            spec = table_specs[tbl_key]
            art = table_artifacts.get(tbl_key)
            if art is None:
                document.add_paragraph(f"{spec.title}: Not available in this run/stage.")
                continue
            document.add_paragraph(f"{spec.title}. {spec.description}")
            df = pd.read_csv(art.csv)
            _docx_add_table(document, df)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _pdf_add_text_page(pdf: PdfPages, *, title: str, paragraphs: Iterable[str]) -> None:
    wrapped_blocks: List[List[str]] = []
    for para in paragraphs:
        text = str(para).strip()
        if not text:
            wrapped_blocks.append([""])
            continue
        block_lines: List[str] = []
        for raw_line in text.splitlines():
            content = raw_line.strip()
            if not content:
                block_lines.append("")
                continue
            block_lines.extend(textwrap.wrap(content, width=118, break_long_words=False) or [""])
        wrapped_blocks.append(block_lines or [""])

    if not wrapped_blocks:
        wrapped_blocks = [["No content."]]

    page_idx = 0
    fig: plt.Figure | None = None
    y = 0.0

    def _start_page() -> None:
        nonlocal fig, y, page_idx
        page_idx += 1
        page_title = title if page_idx == 1 else f"{title} (continued)"
        fig = plt.figure(figsize=(11, 8.5), facecolor="white")
        fig.text(0.06, 0.94, page_title, fontsize=16, weight="bold", va="top")
        y = 0.88

    _start_page()
    for block in wrapped_blocks:
        needed = (len(block) * 0.03) + 0.018
        if y - needed < 0.07:
            assert fig is not None
            pdf.savefig(fig, dpi=REPORT_PDF_DPI)
            plt.close(fig)
            _start_page()
        for line in block:
            assert fig is not None
            fig.text(0.06, y, line, fontsize=11, va="top")
            y -= 0.03
        y -= 0.018

    assert fig is not None
    pdf.savefig(fig, dpi=REPORT_PDF_DPI)
    plt.close(fig)


def _pdf_add_image_page(pdf: PdfPages, *, title: str, caption: str, image_path: Path) -> None:
    fig = plt.figure(figsize=(11, 8.5), facecolor="white")
    fig.text(0.05, 0.965, title, fontsize=14, weight="bold", va="top")
    ax = fig.add_axes([0.05, 0.19, 0.90, 0.72])
    ax.axis("off")
    img = plt.imread(image_path)
    ax.imshow(img)
    wrapped_caption = "\n".join(textwrap.wrap(caption, width=150, break_long_words=False))
    fig.text(0.05, 0.05, wrapped_caption, fontsize=9, va="bottom")
    pdf.savefig(fig, dpi=REPORT_PDF_DPI)
    plt.close(fig)


def _pdf_add_table_pages(pdf: PdfPages, *, title: str, df: pd.DataFrame, max_rows: int = 24) -> None:
    if df.empty:
        _pdf_add_text_page(pdf, title=title, paragraphs=["No table rows available."])
        return

    start = 0
    page_idx = 1
    while start < len(df):
        chunk = df.iloc[start : start + max_rows]
        fig, ax = plt.subplots(figsize=(11, 8.5))
        ax.axis("off")
        subtitle = f"{title} (page {page_idx})" if len(df) > max_rows else title
        ax.set_title(subtitle, fontsize=12, pad=12)
        table = ax.table(cellText=chunk.values, colLabels=chunk.columns, cellLoc="left", loc="center")
        table.auto_set_font_size(False)
        table.set_fontsize(8)
        table.scale(1, 1.2)
        pdf.savefig(fig, dpi=REPORT_PDF_DPI)
        plt.close(fig)
        start += max_rows
        page_idx += 1


def _render_report_pdf(
    *,
    report_title: str,
    sections: Sequence[ReportSection],
    figure_specs: Mapping[str, FigureSpecEnhanced],
    figure_artifacts: Mapping[str, FigureArtifactSet],
    table_specs: Mapping[str, TableSpecEnhanced],
    table_artifacts: Mapping[str, TableArtifactSet],
    output_path: Path,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with PdfPages(output_path) as pdf:
        _pdf_add_text_page(
            pdf,
            title=report_title,
            paragraphs=[
                f"Generated: {datetime.now().isoformat(timespec='seconds')}",
                "Enhanced fairness manuscript artifact report.",
            ],
        )

        for section in sections:
            _pdf_add_text_page(pdf, title=section.title, paragraphs=section.paragraphs)

            for fig_key in section.figure_keys:
                spec = figure_specs[fig_key]
                art = figure_artifacts.get(fig_key)
                if art is None:
                    _pdf_add_text_page(pdf, title=spec.title, paragraphs=["Not available in this run/stage."])
                    continue
                _pdf_add_image_page(pdf, title=spec.title, caption=spec.caption, image_path=art.png)

            for tbl_key in section.table_keys:
                spec = table_specs[tbl_key]
                art = table_artifacts.get(tbl_key)
                if art is None:
                    _pdf_add_text_page(pdf, title=spec.title, paragraphs=["Not available in this run/stage."])
                    continue
                df = pd.read_csv(art.csv)
                _pdf_add_table_pages(pdf, title=spec.title, df=df)


def generate_fairness_enhanced_artifacts(
    *,
    fairness_enhanced_dir: Path,
    output_dir: Path = DEFAULT_ARTIFACTS_DIR,
    seed: int = 0,
    include_supplement: bool = True,
    report_title: str = "Enhanced Fairness Manuscript Report",
    report_stem: str = "enhanced_fairness_report",
) -> EnhancedArtifactBundle:
    np.random.seed(seed)

    fairness_enhanced_dir = fairness_enhanced_dir.expanduser().resolve()
    output_dir = output_dir.expanduser().resolve()
    figures_dir = output_dir / "figures"
    tables_dir = output_dir / "tables"
    output_dir.mkdir(parents=True, exist_ok=True)

    loaded = _load_enhanced_inputs(fairness_enhanced_dir)
    primary_df = loaded["primary"]
    effects_df = loaded["effects"]
    endpoint_payload = loaded["effects_payload"]
    trend_success_df = loaded["trend_success"]
    trend_iou_summary = loaded["trend_iou_summary"]
    threshold_df = loaded["threshold"]
    run_metadata = loaded["run_metadata"]
    run_summary = loaded["run_summary"]
    optional = loaded["optional"]
    covadj_df = optional.get("covadj_effects")
    covadj_payload = optional.get("covadj_payload")
    covadj_components_df = optional.get("covadj_components")

    label_text = endpoint_payload.get("label_text", {}) if isinstance(endpoint_payload, dict) else {}
    proxy_caption = str(label_text.get("figure_caption_snippet", f"{_PROXY_FALLBACK}; {_BINARY_FALLBACK}."))

    figure_specs = _figure_specs(proxy_caption)
    table_specs = _table_specs()

    figures: Dict[str, FigureArtifactSet] = {}
    tables: Dict[str, TableArtifactSet] = {}

    fig_e1 = _render_cohort_accountability(run_summary)
    figures["E1"] = _save_figure_triplet(fig_e1, figure_specs["E1"].stem, figures_dir)
    plt.close(fig_e1)

    fig_e2 = _render_primary_effects_forest(effects_df)
    figures["E2"] = _save_figure_triplet(fig_e2, figure_specs["E2"].stem, figures_dir)
    plt.close(fig_e2)

    fig_e3 = _render_trends_success_iou(
        trend_success_df,
        trend_iou_summary,
        iou_png_path=optional.get("ita_trend_iou_png"),
        covariates_used=run_summary.trend_covariates_used,
    )
    figures["E3"] = _save_figure_triplet(fig_e3, figure_specs["E3"].stem, figures_dir)
    plt.close(fig_e3)

    fig_e4 = _render_threshold_sensitivity(threshold_df)
    figures["E4"] = _save_figure_triplet(fig_e4, figure_specs["E4"].stem, figures_dir)
    plt.close(fig_e4)

    fig_e5 = _render_ita_distribution(primary_df, run_summary.ita_cutoff)
    figures["E5"] = _save_figure_triplet(fig_e5, figure_specs["E5"].stem, figures_dir)
    plt.close(fig_e5)

    table_e1_df = _table_cohort_and_provenance(primary_df, run_summary, optional.get("dedup_report"))
    tables["E1"] = _write_table_artifacts(table_e1_df, spec=table_specs["E1"], output_dir=tables_dir)

    table_e2_df = _table_primary_effects(effects_df)
    tables["E2"] = _write_table_artifacts(table_e2_df, spec=table_specs["E2"], output_dir=tables_dir)

    table_e3_df = _table_group_descriptives(primary_df)
    tables["E3"] = _write_table_artifacts(table_e3_df, spec=table_specs["E3"], output_dir=tables_dir)

    table_e4_df = _table_trend_model_spec(run_metadata, run_summary)
    tables["E4"] = _write_table_artifacts(table_e4_df, spec=table_specs["E4"], output_dir=tables_dir)
    if isinstance(covadj_df, pd.DataFrame) and not covadj_df.empty:
        table_e5_df = _table_covadj_effects(covadj_df, covadj_payload if isinstance(covadj_payload, dict) else None)
        tables["E5"] = _write_table_artifacts(table_e5_df, spec=table_specs["E5"], output_dir=tables_dir)
    if isinstance(covadj_components_df, pd.DataFrame) and not covadj_components_df.empty:
        table_e6_df = _table_covadj_components(covadj_components_df)
        tables["E6"] = _write_table_artifacts(table_e6_df, spec=table_specs["E6"], output_dir=tables_dir)

    if include_supplement:
        dedup_df = optional.get("dedup_sensitivity")
        if isinstance(dedup_df, pd.DataFrame):
            fig_es1 = _render_dedup_sensitivity(dedup_df)
            figures["ES1"] = _save_figure_triplet(fig_es1, figure_specs["ES1"].stem, figures_dir)
            plt.close(fig_es1)
            tables["ES2"] = _write_table_artifacts(dedup_df, spec=table_specs["ES2"], output_dir=tables_dir)

        mask_df = optional.get("mask_source_sensitivity")
        if isinstance(mask_df, pd.DataFrame):
            fig_es2 = _render_mask_source_sensitivity(mask_df)
            figures["ES2"] = _save_figure_triplet(fig_es2, figure_specs["ES2"].stem, figures_dir)
            plt.close(fig_es2)
            tables["ES3"] = _write_table_artifacts(mask_df, spec=table_specs["ES3"], output_dir=tables_dir)

        qc_fig = _render_covariate_qc_overview_from_frame(primary_df)
        if qc_fig is None:
            qc_fig = _render_covariate_qc_overview(optional.get("covariates_qc_dir"))
        if qc_fig is not None:
            figures["ES3"] = _save_figure_triplet(qc_fig, figure_specs["ES3"].stem, figures_dir)
            plt.close(qc_fig)

        tables["ES1"] = _write_table_artifacts(threshold_df, spec=table_specs["ES1"], output_dir=tables_dir)

        runtime_profile_table = _table_runtime_profile(run_metadata, optional.get("runtime_profile"))
        tables["ES4"] = _write_table_artifacts(runtime_profile_table, spec=table_specs["ES4"], output_dir=tables_dir)

    available_supp_figs = sorted([key for key in figures if key.startswith("ES")])
    available_supp_tables = sorted([key for key in tables if key.startswith("ES")])

    sections = _build_sections(
        report_title=report_title,
        run_summary=run_summary,
        run_metadata=run_metadata,
        endpoint_payload=endpoint_payload,
        effects_df=effects_df,
        covadj_df=covadj_df if isinstance(covadj_df, pd.DataFrame) else None,
        covadj_components_df=covadj_components_df if isinstance(covadj_components_df, pd.DataFrame) else None,
        covadj_payload=covadj_payload if isinstance(covadj_payload, dict) else None,
        include_supplement=include_supplement,
        available_supp_figures=available_supp_figs,
        available_supp_tables=available_supp_tables,
    )

    report_md = output_dir / f"{report_stem}.md"
    report_html = output_dir / f"{report_stem}.html"
    report_docx = output_dir / f"{report_stem}.docx"
    report_pdf = output_dir / f"{report_stem}.pdf"

    _render_report_markdown(
        report_title=report_title,
        sections=sections,
        figure_specs=figure_specs,
        figure_artifacts=figures,
        table_specs=table_specs,
        table_artifacts=tables,
        output_path=report_md,
    )
    _render_report_html(
        report_title=report_title,
        sections=sections,
        figure_specs=figure_specs,
        figure_artifacts=figures,
        table_specs=table_specs,
        table_artifacts=tables,
        output_path=report_html,
    )
    _render_report_docx(
        report_title=report_title,
        sections=sections,
        figure_specs=figure_specs,
        figure_artifacts=figures,
        table_specs=table_specs,
        table_artifacts=tables,
        output_path=report_docx,
    )
    _render_report_pdf(
        report_title=report_title,
        sections=sections,
        figure_specs=figure_specs,
        figure_artifacts=figures,
        table_specs=table_specs,
        table_artifacts=tables,
        output_path=report_pdf,
    )

    LOGGER.info("Enhanced fairness manuscript artifacts written under %s", output_dir)
    return EnhancedArtifactBundle(
        figures=figures,
        tables=tables,
        report_pdf=report_pdf,
        report_docx=report_docx,
        report_markdown=report_md,
        report_html=report_html,
    )


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render enhanced fairness manuscript artifacts.")
    parser.add_argument(
        "--fairness-enhanced-dir",
        type=Path,
        required=True,
        help="Directory containing fairness_enhanced artifacts from a completed enhanced fairness run.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_ARTIFACTS_DIR,
        help="Destination directory for generated figures/tables/reports.",
    )
    parser.add_argument(
        "--seed",
        type=int,
        default=0,
        help="Deterministic seed for plot-level stochastic elements.",
    )
    parser.add_argument(
        "--include-supplement",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Include supplementary sensitivity/QC artifacts when available.",
    )
    parser.add_argument(
        "--report-title",
        default="Enhanced Fairness Manuscript Report",
        help="Title shown in report outputs.",
    )
    parser.add_argument(
        "--report-stem",
        default="enhanced_fairness_report",
        help="Base filename stem for report outputs (md/html/pdf/docx).",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(name)s:%(message)s")
    generate_fairness_enhanced_artifacts(
        fairness_enhanced_dir=args.fairness_enhanced_dir,
        output_dir=args.output_dir,
        seed=args.seed,
        include_supplement=bool(args.include_supplement),
        report_title=str(args.report_title),
        report_stem=str(args.report_stem),
    )


if __name__ == "__main__":
    main()

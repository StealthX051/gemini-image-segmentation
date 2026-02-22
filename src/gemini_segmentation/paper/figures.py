from __future__ import annotations

"""Render fairness figures and tables for the manuscript."""

import argparse
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Mapping, Optional, Sequence, Tuple

import matplotlib

# Use a non-interactive backend so figure generation is reliable in headless runs.
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from matplotlib.colors import to_rgb
from scipy import stats as scipy_stats
from scipy.stats import bootstrap

from ..fairness import compute_fairness_statistics, summarize_groups
from ..types import FairnessResult, GroupSummary

LOGGER = logging.getLogger(__name__)

DEFAULT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_ARTIFACTS_DIR = DEFAULT_ROOT / "artifacts" / "fairness"

BINARY_TONE_PALETTE = {"Light": "#377eb8", "Dark": "#ff7f0e"}
PANEL_FACE = "#f6f7fb"
GRID_COLOR = "#8f8f8f"


@dataclass
class FigureArtifacts:
    pdf: Path
    png: Path
    svg: Path


@dataclass
class TableArtifacts:
    csv: Path
    html: Path
    docx: Path


def _save_figure2_outputs(fig: plt.Figure, output_dir: Path) -> FigureArtifacts:
    """Save Figure 2 in raster and vector formats."""

    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / "figure2.pdf"
    png_path = output_dir / "figure2.png"
    svg_path = output_dir / "figure2.svg"
    fig.savefig(pdf_path)
    fig.savefig(png_path)
    fig.savefig(svg_path)
    return FigureArtifacts(pdf=pdf_path, png=png_path, svg=svg_path)


def _results_to_df(results: Iterable[FairnessResult]) -> pd.DataFrame:
    """Convert :class:`FairnessResult` records into a DataFrame."""

    return pd.DataFrame([r.__dict__ for r in results])


def _summaries_to_df(summaries: Iterable[GroupSummary]) -> pd.DataFrame:
    """Convert :class:`GroupSummary` records into a DataFrame."""

    rows = []
    for s in summaries:
        rows.append(
            {
                "group": s.group_name,
                "count": s.count,
                "mean_iou": s.mean_iou,
                "median_iou": s.median_iou,
                "ci_iou_lower": s.ci_iou.lower,
                "ci_iou_upper": s.ci_iou.upper,
                "mean_dice": s.mean_dice,
                "median_dice": s.median_dice,
                "ci_dice_lower": s.ci_dice.lower,
                "ci_dice_upper": s.ci_dice.upper,
                "success_rate": s.success_rate,
            }
        )
    return pd.DataFrame(rows)


def _load_csv(path: Path, description: str) -> pd.DataFrame:
    """Load a CSV file with a helpful error if it is missing."""

    if not path.exists():
        raise FileNotFoundError(f"{description} not found: {path}")
    return pd.read_csv(path)


def _load_fairness_dir(fairness_dir: Path) -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, float]]:
    """Load fairness CSV outputs from a run directory."""

    results_path = fairness_dir / "fairness_results.csv"
    summary_path = fairness_dir / "fairness_summary.csv"
    stats_path = fairness_dir / "fairness_stats.csv"

    results_df = _load_csv(results_path, "Fairness results")
    summary_df = _load_csv(summary_path, "Fairness summary")
    stats_df = _load_csv(stats_path, "Fairness statistics") if stats_path.exists() else pd.DataFrame()
    stats_payload: Dict[str, float] = stats_df.iloc[0].to_dict() if not stats_df.empty else {}
    return results_df, summary_df, stats_payload


def _tone_order(df: pd.DataFrame) -> List[str]:
    order = ["Light", "Dark"]
    present = [tone for tone in order if tone in df["tone_group"].unique().tolist()]
    remaining = [tone for tone in df["tone_group"].unique().tolist() if tone not in present]
    return present + remaining


def _adjust_lightness(color: str, factor: float = 1.2) -> Tuple[float, float, float]:
    rgb = np.array(to_rgb(color), dtype=float)
    if factor >= 1.0:
        return tuple(np.clip(rgb + (1.0 - rgb) * (factor - 1.0), 0.0, 1.0))
    return tuple(np.clip(rgb * factor, 0.0, 1.0))


def _format_p_value(p_val: float | None) -> str:
    if p_val is None or pd.isna(p_val):
        return "n/a"
    if p_val < 0.001:
        return "<0.001"
    return f"={p_val:.3f}"


def _p_to_stars(p_val: float | None) -> str:
    if p_val is None or pd.isna(p_val):
        return "n.s."
    if p_val < 0.001:
        return "***"
    if p_val < 0.01:
        return "**"
    if p_val < 0.05:
        return "*"
    return "n.s."


def _cliffs_delta(x: np.ndarray, y: np.ndarray) -> float:
    clean_x = np.asarray([v for v in x if not pd.isna(v)], dtype=float)
    clean_y = np.asarray([v for v in y if not pd.isna(v)], dtype=float)
    if clean_x.size == 0 or clean_y.size == 0:
        return float("nan")
    diffs = clean_x[:, None] - clean_y[None, :]
    wins = np.sum(diffs > 0)
    losses = np.sum(diffs < 0)
    return float((wins - losses) / (clean_x.size * clean_y.size))


def _two_group_p_value(df: pd.DataFrame) -> float:
    order = _tone_order(df)
    if len(order) < 2:
        return float("nan")
    first = df[df["tone_group"] == order[0]]["iou"].dropna().to_numpy()
    second = df[df["tone_group"] == order[1]]["iou"].dropna().to_numpy()
    if first.size < 2 or second.size < 2:
        return float("nan")
    try:
        return float(scipy_stats.kruskal(first, second).pvalue)
    except Exception:
        return float("nan")


def _render_histogram(ax: plt.Axes, df: pd.DataFrame) -> None:
    ax.set_facecolor(PANEL_FACE)
    ita = df["ita"].dropna()
    ax.hist(ita, bins=40, color="#4c72b0", edgecolor="white", linewidth=0.6, alpha=0.9)
    ax.axvline(28, color="red", linestyle="--", label="28° threshold")
    ax.text(28.5, ax.get_ylim()[1] * 0.9, "28° cut-point", color="red", fontsize=8, va="top")
    ax.set_title("ITA distribution")
    ax.set_xlabel("ITA (°)")
    ax.set_ylabel("Count")
    ax.grid(axis="y", linestyle=":", color=GRID_COLOR, alpha=0.35)
    ax.legend()


def _render_success_bars(ax: plt.Axes, df: pd.DataFrame, stats_payload: Dict[str, float]) -> None:
    ax.set_facecolor(PANEL_FACE)
    order = _tone_order(df)
    counts = df.groupby(["tone_group", "success"]).size().unstack(fill_value=0)
    success = counts.get(True, pd.Series([0] * len(order), index=order)).reindex(order, fill_value=0)
    failure = counts.get(False, pd.Series([0] * len(order), index=order)).reindex(order, fill_value=0)
    total = (success + failure).replace(0, np.nan)
    success_prop = (success / total).fillna(0.0)
    failure_prop = (failure / total).fillna(0.0)

    positions = np.arange(len(order))
    failure_colors = [_adjust_lightness(BINARY_TONE_PALETTE.get(tone, "#7f7f7f"), 0.8) for tone in order]
    success_colors = [_adjust_lightness(BINARY_TONE_PALETTE.get(tone, "#7f7f7f"), 1.15) for tone in order]

    bar_height = 0.56
    bars_fail = ax.barh(
        positions,
        failure_prop.values,
        color=failure_colors,
        edgecolor="white",
        label="Fail",
        height=bar_height,
    )
    bars_ok = ax.barh(
        positions,
        success_prop.values,
        left=failure_prop.values,
        color=success_colors,
        edgecolor="white",
        label="Success",
        height=bar_height,
    )

    for i, (bar_fail, bar_ok) in enumerate(zip(bars_fail, bars_ok)):
        fail_count = int(failure.iloc[i])
        succ_count = int(success.iloc[i])

        fail_width = float(bar_fail.get_width())
        fail_text = f"Fail\n{fail_count}"
        fail_y = bar_fail.get_y() + bar_fail.get_height() / 2.0
        if fail_width >= 0.12:
            fail_x = bar_fail.get_x() + fail_width / 2.0
            fail_ha = "center"
            fail_color = "white"
        else:
            fail_x = bar_fail.get_x() + fail_width + 0.01
            fail_ha = "left"
            fail_color = "#303030"
        ax.text(
            fail_x,
            fail_y,
            fail_text,
            ha=fail_ha,
            va="center",
            color=fail_color,
            fontsize=8,
            fontweight="bold",
            linespacing=0.9,
        )

        succ_width = float(bar_ok.get_width())
        succ_text = f"Success\n{succ_count}"
        succ_y = bar_ok.get_y() + bar_ok.get_height() / 2.0
        if succ_width >= 0.15:
            succ_x = bar_ok.get_x() + succ_width / 2.0
            succ_ha = "center"
            succ_color = "white"
        else:
            succ_x = bar_ok.get_x() + succ_width + 0.01
            succ_ha = "left"
            succ_color = "#303030"
        ax.text(
            succ_x,
            succ_y,
            succ_text,
            ha=succ_ha,
            va="center",
            color=succ_color,
            fontsize=8,
            fontweight="bold",
            linespacing=0.9,
        )

    tick_labels = [f"{tone} (n={int(total.iloc[idx])})" for idx, tone in enumerate(order)]
    ax.set_yticks(positions)
    ax.set_yticklabels(tick_labels)
    ax.set_xlim(0, 1.0)
    ax.xaxis.set_major_formatter(matplotlib.ticker.PercentFormatter(1.0))
    ax.set_xlabel("Proportion of Images")
    ax.set_title("Success Rate (IoU ≥ 0.5)")
    ax.grid(axis="x", linestyle=":", alpha=0.35)
    ax.set_ylim(-0.5, len(order) - 0.5)

    if {"Light", "Dark"}.issubset(set(order)):
        light_idx = order.index("Light")
        dark_idx = order.index("Dark")
        delta = (success_prop.iloc[light_idx] - success_prop.iloc[dark_idx]) * 100.0
        chi2_p = stats_payload.get("chi2_success_p", np.nan)
        cliff_delta = stats_payload.get("cliffs_delta_iou_light_dark", np.nan)
        star = _p_to_stars(chi2_p)
        cliff_text = ""
        if cliff_delta is not None and not pd.isna(cliff_delta):
            cliff_text = f"\nCliff's d {cliff_delta:+.2f}"
        ax.text(
            1.02,
            0.95,
            f"Delta = {delta:+.1f}%\nchi2 p {_format_p_value(chi2_p)} ({star}){cliff_text}",
            transform=ax.transAxes,
            ha="left",
            va="top",
            fontsize=9,
            color="#333333",
        )


def _render_distribution(
    ax: plt.Axes,
    df: pd.DataFrame,
    title: str,
    *,
    p_value: float | None = None,
    cliff_delta: float | None = None,
    y_min: float = 0.0,
    success_panel: bool = False,
) -> None:
    ax.set_facecolor(PANEL_FACE)
    order = _tone_order(df)
    data = [df[df["tone_group"] == tone]["iou"].dropna().values for tone in order]
    positions = np.arange(len(order)) + 1

    if not any(len(values) for values in data):
        ax.text(0.5, 0.5, "No IoU data", ha="center", va="center", transform=ax.transAxes)
        ax.set_xticks(positions)
        ax.set_xticklabels(order)
        ax.set_ylim(y_min, 1.05)
        ax.set_title(title)
        return

    violin = ax.violinplot(
        data,
        positions=positions,
        showmeans=False,
        showmedians=False,
        showextrema=False,
        widths=0.85,
        bw_method="scott",
    )
    max_n = max((len(values) for values in data), default=0)
    for idx, body in enumerate(violin["bodies"]):
        tone = order[idx] if idx < len(order) else ""
        body.set_facecolor(BINARY_TONE_PALETTE.get(tone, "#4c72b0"))
        body.set_alpha(0.72)
        body.set_edgecolor("#202020")
        body.set_linewidth(0.9)
        if max_n > 0 and len(data[idx]) > 0:
            scale = np.sqrt(len(data[idx])) / np.sqrt(max_n)
            verts = body.get_paths()[0].vertices
            center = float(np.mean(verts[:, 0]))
            verts[:, 0] = (verts[:, 0] - center) * scale + center

    ax.boxplot(
        data,
        positions=positions,
        widths=0.15,
        patch_artist=True,
        showfliers=False,
        medianprops={"color": "#e41a1c", "linewidth": 1.2},
        boxprops={"linewidth": 1.2, "facecolor": "white"},
        whiskerprops={"linewidth": 1.1},
        capprops={"linewidth": 1.1},
    )
    ax.set_xticks(positions)
    ax.set_xticklabels(order)
    ax.set_xlabel("Skin tone group")
    ax.set_ylabel("Intersection over Union (IoU)")
    ax.set_title(title)
    ax.set_ylim(y_min, 1.04)
    ax.grid(axis="y", linestyle=":", color=GRID_COLOR, alpha=0.35)

    if success_panel:
        ax.axhline(0.5, linestyle="-", linewidth=1.2, color="#d62728", zorder=0)
    elif y_min <= 0.5:
        ax.axhline(0.5, linestyle="--", linewidth=1.1, color="#7a7a7a", zorder=0)
    ax.axhline(0.8, linestyle=":", linewidth=1.0, color="#7a7a7a", zorder=0)

    y_top = ax.get_ylim()[1]
    for pos, values in zip(positions, data):
        if len(values) == 0:
            continue
        median = float(np.median(values))
        ax.text(
            pos,
            y_top - 0.005,
            f"n={len(values)}\nmedian={median:.2f}",
            ha="center",
            va="top",
            fontsize=8,
            bbox={"facecolor": "white", "alpha": 0.75, "edgecolor": "none", "pad": 0.2},
        )

    if len(order) == 2 and p_value is not None and not pd.isna(p_value):
        y_bracket = y_top - 0.045
        h = 0.006
        x1, x2 = positions[0], positions[1]
        ax.plot(
            [x1, x1, x2, x2],
            [y_bracket, y_bracket + h, y_bracket + h, y_bracket],
            lw=1.2,
            color="#111111",
            clip_on=False,
        )
        ax.text(
            (x1 + x2) * 0.5,
            y_bracket + h + 0.001,
            f"{_p_to_stars(p_value)}  (p {_format_p_value(p_value)})",
            ha="center",
            va="bottom",
            fontsize=10,
            color="#111111",
            fontweight="bold",
        )
        if cliff_delta is not None and not pd.isna(cliff_delta):
            ax.text(
                (x1 + x2) * 0.5,
                y_bracket - 0.006,
                f"Cliff's d={cliff_delta:+.2f}",
                ha="center",
                va="top",
                fontsize=8,
                color="#333333",
            )

    if y_min > 0.0:
        ax.text(
            0.01,
            0.02,
            f"Scale truncated at {y_min:.1f}",
            transform=ax.transAxes,
            fontsize=8,
            color="#555555",
        )


def render_figure2(
    df: pd.DataFrame,
    *,
    output_dir: Path,
    stats_payload: Optional[Dict[str, float]] = None,
    seed: int = 0,
) -> FigureArtifacts:
    """Render the four-panel fairness figure and write PNG/PDF/SVG outputs."""

    if df.empty:
        raise ValueError("Fairness results are empty; cannot render figure.")
    np.random.seed(seed)

    fig, axes = plt.subplots(2, 2, figsize=(13, 9))
    fig.patch.set_facecolor("white")
    stats_payload = stats_payload or {}
    _render_histogram(axes[0, 0], df)
    _render_success_bars(axes[0, 1], df, stats_payload)

    all_p = stats_payload.get("dunn_iou_light_dark_p", stats_payload.get("kruskal_iou_p", np.nan))
    all_delta = stats_payload.get("cliffs_delta_iou_light_dark", np.nan)
    _render_distribution(
        axes[1, 0],
        df,
        "IoU by tone (all)",
        p_value=all_p,
        cliff_delta=all_delta,
        y_min=0.0,
        success_panel=False,
    )

    filtered = df[df["iou"] >= 0.5]
    filtered_p = _two_group_p_value(filtered)
    order = _tone_order(filtered)
    if len(order) >= 2:
        arr_a = filtered[filtered["tone_group"] == order[0]]["iou"].dropna().to_numpy()
        arr_b = filtered[filtered["tone_group"] == order[1]]["iou"].dropna().to_numpy()
        filtered_delta = _cliffs_delta(arr_a, arr_b)
    else:
        filtered_delta = float("nan")
    _render_distribution(
        axes[1, 1],
        filtered,
        "IoU by tone (IoU ≥ 0.5)",
        p_value=filtered_p,
        cliff_delta=filtered_delta,
        y_min=0.5,
        success_panel=True,
    )
    plt.tight_layout()

    artifacts = _save_figure2_outputs(fig, output_dir)
    plt.close(fig)
    LOGGER.info("Figure 2 saved to %s", output_dir)
    return artifacts


def _format_ci(mean: float, lower: float, upper: float) -> str:
    return f"{mean:.3f} ({lower:.3f}, {upper:.3f})"


def _bootstrap_ci(
    values: Sequence[float],
    statistic: Callable[[np.ndarray], float],
    *,
    seed: int = 0,
    n_resamples: int = 5000,
    method: str = "bca",
) -> Tuple[float, float]:
    """Compute a BCa confidence interval for a statistic with deterministic sampling."""

    clean = np.asarray([v for v in values if not pd.isna(v)], dtype=float)
    if clean.size < 2:
        return float("nan"), float("nan")

    try:
        res = bootstrap(
            (clean,),
            statistic,
            vectorized=False,
            method=method,
            n_resamples=n_resamples,
            confidence_level=0.95,
            random_state=np.random.default_rng(seed),
        )
        return float(res.confidence_interval.low), float(res.confidence_interval.high)
    except Exception as exc:  # pragma: no cover - fallback path
        LOGGER.warning("Falling back to percentile CI for %s: %s", statistic.__name__, exc)
        rng = np.random.default_rng(seed)
        samples = [statistic(rng.choice(clean, size=clean.size, replace=True)) for _ in range(n_resamples)]
        return float(np.percentile(samples, 2.5)), float(np.percentile(samples, 97.5))


def _group_stats(results_df: pd.DataFrame, group: str, *, seed: int = 0) -> Mapping[str, Tuple[float, float, float]]:
    """Compute mean/median IoU and Dice metrics with BCa CIs for a tone group."""

    group_df = results_df[results_df["tone_group"] == group]
    if group_df.empty:
        raise ValueError(f"No data available for group '{group}'")

    ious = group_df["iou"].dropna().to_numpy()
    dices = group_df["dice"].dropna().to_numpy()

    stats: Dict[str, Tuple[float, float, float]] = {
        "mean_iou": (float(np.mean(ious)) if ious.size else float("nan"),)
        + _bootstrap_ci(ious, np.mean, seed=seed),
        "median_iou": (float(np.median(ious)) if ious.size else float("nan"),)
        + _bootstrap_ci(ious, np.median, seed=seed),
        "mean_dice": (float(np.mean(dices)) if dices.size else float("nan"),)
        + _bootstrap_ci(dices, np.mean, seed=seed),
        "median_dice": (float(np.median(dices)) if dices.size else float("nan"),)
        + _bootstrap_ci(dices, np.median, seed=seed),
    }
    return stats


def render_table4(
    *,
    results_df: pd.DataFrame,
    summary_df: pd.DataFrame,
    stats_payload: Dict[str, float],
    output_dir: Path,
    seed: int = 0,
) -> TableArtifacts:
    """Render Table 4 (summary statistics and tests) to CSV/HTML/DOCX."""

    if results_df.empty:
        raise ValueError("Fairness results are empty; cannot render table.")

    groups = set(results_df["tone_group"].unique())
    if not {"Light", "Dark"}.issubset(groups):
        raise ValueError("Both Light and Dark groups are required for Table 4.")

    light_stats = _group_stats(results_df, "Light", seed=seed)
    dark_stats = _group_stats(results_df, "Dark", seed=seed)

    rows = []
    rows.append(
        {
            "metric": "IoU mean",
            "Light": _format_ci(*light_stats["mean_iou"]),
            "Dark": _format_ci(*dark_stats["mean_iou"]),
            "kruskal_p": stats_payload.get("kruskal_iou_p", np.nan),
            "cliffs_delta": stats_payload.get("cliffs_delta_iou_light_dark", np.nan),
            "chi2": stats_payload.get("chi2_success", np.nan),
            "chi2_p": stats_payload.get("chi2_success_p", np.nan),
        }
    )
    rows.append(
        {
            "metric": "IoU median",
            "Light": _format_ci(*light_stats["median_iou"]),
            "Dark": _format_ci(*dark_stats["median_iou"]),
            "kruskal_p": stats_payload.get("kruskal_iou_p", np.nan),
            "cliffs_delta": stats_payload.get("cliffs_delta_iou_light_dark", np.nan),
            "chi2": stats_payload.get("chi2_success", np.nan),
            "chi2_p": stats_payload.get("chi2_success_p", np.nan),
        }
    )
    rows.append(
        {
            "metric": "Dice mean",
            "Light": _format_ci(*light_stats["mean_dice"]),
            "Dark": _format_ci(*dark_stats["mean_dice"]),
            "kruskal_p": stats_payload.get("kruskal_dice_p", np.nan),
            "cliffs_delta": np.nan,
            "chi2": stats_payload.get("chi2_success", np.nan),
            "chi2_p": stats_payload.get("chi2_success_p", np.nan),
        }
    )
    rows.append(
        {
            "metric": "Dice median",
            "Light": _format_ci(*light_stats["median_dice"]),
            "Dark": _format_ci(*dark_stats["median_dice"]),
            "kruskal_p": stats_payload.get("kruskal_dice_p", np.nan),
            "cliffs_delta": np.nan,
            "chi2": stats_payload.get("chi2_success", np.nan),
            "chi2_p": stats_payload.get("chi2_success_p", np.nan),
        }
    )

    table_df = pd.DataFrame(rows)
    output_dir.mkdir(parents=True, exist_ok=True)
    csv_path = output_dir / "table4.csv"
    html_path = output_dir / "table4.html"
    docx_path = output_dir / "table4.docx"

    table_df.to_csv(csv_path, index=False)
    table_df.to_html(html_path, index=False)

    document = Document()
    document.add_heading("Table 4: Fairness metrics", 0)
    table = document.add_table(rows=1, cols=len(table_df.columns))
    for idx, col in enumerate(table_df.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    document.save(docx_path)
    LOGGER.info("Table 4 saved to %s", output_dir)
    return TableArtifacts(csv=csv_path, html=html_path, docx=docx_path)


def generate_fairness_artifacts(
    *,
    fairness_dir: Optional[Path] = None,
    results: Optional[Sequence[FairnessResult]] = None,
    summaries: Optional[Sequence[GroupSummary]] = None,
    stats_payload: Optional[Dict[str, float]] = None,
    output_dir: Path = DEFAULT_ARTIFACTS_DIR,
    seed: int = 0,
) -> Tuple[FigureArtifacts, TableArtifacts]:
    """Build Figure 2 and Table 4 from fairness CSVs or in-memory results."""

    if fairness_dir is None and results is None:
        raise ValueError("Provide either fairness_dir or in-memory results.")

    if fairness_dir:
        LOGGER.info("Loading fairness artifacts from %s", fairness_dir)
        results_df, summary_df, stats_payload_from_disk = _load_fairness_dir(fairness_dir)
        stats_payload = stats_payload or stats_payload_from_disk
    else:
        if results is None:
            raise ValueError("Results are required when fairness_dir is not provided.")
        if summaries is None:
            summaries = summarize_groups(results)
        if stats_payload is None:
            stats_payload = compute_fairness_statistics(list(results), 0.5)
        results_df = _results_to_df(results)
        summary_df = _summaries_to_df(summaries)

    stats_payload = stats_payload or {}
    figure_artifacts = render_figure2(results_df, output_dir=output_dir, stats_payload=stats_payload, seed=seed)
    table_artifacts = render_table4(
        results_df=results_df,
        summary_df=summary_df,
        stats_payload=stats_payload,
        output_dir=output_dir,
        seed=seed,
    )
    return figure_artifacts, table_artifacts


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render fairness figures and tables")
    parser.add_argument("--fairness-dir", type=Path, required=True, help="Directory containing fairness CSV outputs")
    parser.add_argument(
        "--output-dir", type=Path, default=DEFAULT_ARTIFACTS_DIR, help="Destination for generated artifacts"
    )
    parser.add_argument(
        "--seed", type=int, default=0, help="Seed for deterministic plots (affects violin jitter)",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(name)s:%(message)s")
    generate_fairness_artifacts(
        fairness_dir=args.fairness_dir,
        output_dir=args.output_dir,
        seed=args.seed,
    )


if __name__ == "__main__":
    main()

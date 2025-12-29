from __future__ import annotations

"""Generate paper-ready tables and figures from long-form results."""

import argparse
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document

from ..metrics import calculate_bootstrap_ci
from .config import FigureSpec, PaperConfig, TableSpec, load_paper_config

LOGGER = logging.getLogger(__name__)
DEFAULT_ROOT = Path(__file__).resolve().parents[3]
DEFAULT_CONFIG_PATH = DEFAULT_ROOT / "configs" / "paper.yaml"
DEFAULT_ARTIFACTS_DIR = DEFAULT_ROOT / "artifacts"

COLUMN_LABELS = {
    "task": "Task",
    "model": "Model",
    "prompt_strategy": "Prompt strategy",
    "n": "N",
}


@dataclass
class AggregatedRow:
    values: Dict[str, object]


def _validate_columns(df: pd.DataFrame, required: Iterable[str]) -> None:
    missing = set(required) - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns: {', '.join(sorted(missing))}")


def _load_results(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Results file not found: {path}")
    if path.suffix.lower() == ".csv":
        return pd.read_csv(path)
    if path.suffix.lower() in {".parquet", ".pq"}:
        return pd.read_parquet(path)
    raise ValueError(f"Unsupported results extension: {path.suffix}")


def _display_value(value: object, mapping: Dict[str, str]) -> str:
    if isinstance(value, str):
        return mapping.get(value, value)
    return mapping.get(str(value), str(value))


def _append_display_columns(df: pd.DataFrame, config: PaperConfig) -> pd.DataFrame:
    df = df.copy()
    if "task" in df.columns:
        df["task_label"] = df["task"].apply(lambda v: _display_value(v, config.tasks))
    if "model" in df.columns:
        df["model_label"] = df["model"].apply(lambda v: _display_value(v, config.models))
    if "prompt_strategy" in df.columns:
        df["prompt_label"] = df["prompt_strategy"].apply(
            lambda v: _display_value(v, config.prompt_strategies)
        )
    return df


def _statistic(values: List[float], statistic: str) -> float:
    if statistic == "median":
        return float(np.median(values))
    return float(np.mean(values))


def _aggregate(
    df: pd.DataFrame,
    spec: TableSpec,
    *,
    config: PaperConfig,
) -> pd.DataFrame:
    grouped = df.groupby(spec.group_by)
    rows: List[AggregatedRow] = []
    for key, group in grouped:
        values = group[spec.metric].dropna().tolist()
        if not values:
            continue
        ci = calculate_bootstrap_ci(
            values,
            n_resamples=config.bootstrap_resamples,
            method=config.bootstrap_method,
        )
        stat_value = _statistic(values, spec.statistic)
        if len(spec.group_by) == 1:
            key = (key,)
        row_data: Dict[str, object] = {
            column: key[idx] for idx, column in enumerate(spec.group_by)
        }
        row_data[f"{spec.metric}_{spec.statistic}"] = stat_value
        row_data[f"{spec.metric}_ci_lower"] = ci.lower
        row_data[f"{spec.metric}_ci_upper"] = ci.upper
        row_data["n"] = len(values)
        rows.append(AggregatedRow(values=row_data))
    if not rows:
        raise ValueError(f"No data available for table {spec.name}")
    output = pd.DataFrame([row.values for row in rows])
    return _append_display_columns(output, config)


def _pretty_columns(df: pd.DataFrame) -> pd.DataFrame:
    renamed = {}
    for col in df.columns:
        renamed[col] = COLUMN_LABELS.get(col, col.replace("_", " ").title())
    return df.rename(columns=renamed)


def _write_table_outputs(df: pd.DataFrame, table_dir: Path, name: str, title: str) -> None:
    table_dir.mkdir(parents=True, exist_ok=True)
    csv_path = table_dir / f"{name}.csv"
    html_path = table_dir / f"{name}.html"
    docx_path = table_dir / f"{name}.docx"

    df.to_csv(csv_path, index=False)
    df.to_html(html_path, index=False)

    document = Document()
    document.add_heading(title, 0)
    table = document.add_table(rows=1, cols=len(df.columns))
    for idx, col in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    document.save(docx_path)

    LOGGER.info("Wrote table %s (CSV/HTML/DOCX)", name)


def _render_table(spec: TableSpec, df: pd.DataFrame, config: PaperConfig, table_dir: Path) -> None:
    aggregated = _aggregate(df, spec, config=config)
    ordered_cols = [
        col
        for col in (
            "task_label",
            "model_label",
            "prompt_label",
            f"{spec.metric}_{spec.statistic}",
            f"{spec.metric}_ci_lower",
            f"{spec.metric}_ci_upper",
            "n",
        )
        if col in aggregated.columns
    ]
    display_df = _pretty_columns(aggregated[ordered_cols])
    _write_table_outputs(display_df, table_dir, spec.name, spec.title)


def _render_figure(spec: FigureSpec, df: pd.DataFrame, config: PaperConfig, figure_dir: Path) -> None:
    table_spec = TableSpec(
        name=spec.name,
        title=spec.title,
        metric=spec.metric,
        statistic=spec.statistic,
        group_by=spec.group_by,
    )
    aggregated = _aggregate(df, table_spec, config=config)
    x_col = f"{spec.x}_label" if f"{spec.x}_label" in aggregated.columns else spec.x
    hue_col = f"{spec.hue}_label" if spec.hue and f"{spec.hue}_label" in aggregated.columns else spec.hue
    metric_col = f"{spec.metric}_{spec.statistic}"

    plt.figure(figsize=(8, 4))
    if hue_col:
        categories = aggregated[hue_col].unique()
        x_values = aggregated[x_col].unique()
        width = 0.8 / max(len(categories), 1)
        offsets = np.linspace(-0.4 + width / 2, 0.4 - width / 2, len(categories))
        for offset, category in zip(offsets, categories):
            subset = aggregated[aggregated[hue_col] == category]
            series = []
            for x_value in x_values:
                match = subset[subset[x_col] == x_value]
                if not match.empty:
                    series.append(float(match[metric_col].iloc[0]))
                else:
                    series.append(0.0)
            positions = np.arange(len(x_values)) + offset
            plt.bar(positions, series, width=width, label=str(category))
        plt.xticks(np.arange(len(x_values)), x_values, rotation=30, ha="right")
        plt.legend()
    else:
        plt.bar(aggregated[x_col], aggregated[metric_col], color="#4c72b0")
    plt.title(spec.title)
    plt.ylabel(metric_col.replace("_", " ").title())
    plt.tight_layout()

    figure_dir.mkdir(parents=True, exist_ok=True)
    png_path = figure_dir / f"{spec.name}.png"
    pdf_path = figure_dir / f"{spec.name}.pdf"
    plt.savefig(png_path)
    plt.savefig(pdf_path)
    plt.close()
    LOGGER.info("Wrote figure %s (PNG/PDF)", spec.name)


def generate_artifacts(
    results_path: Path,
    *,
    config_path: Path = DEFAULT_CONFIG_PATH,
    artifacts_dir: Path = DEFAULT_ARTIFACTS_DIR,
) -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(name)s:%(message)s")
    LOGGER.info("Loading paper config from %s", config_path)
    config = load_paper_config(config_path)

    LOGGER.info("Loading results from %s", results_path)
    df = _load_results(results_path)
    _validate_columns(df, config.required_columns)
    df = _append_display_columns(df, config)

    table_dir = artifacts_dir / "tables"
    figure_dir = artifacts_dir / "figures"

    for table in config.tables:
        LOGGER.info("Rendering table %s", table.name)
        _render_table(table, df, config, table_dir)

    for figure in config.figures:
        LOGGER.info("Rendering figure %s", figure.name)
        _render_figure(figure, df, config, figure_dir)

    LOGGER.info("Artifacts ready under %s", artifacts_dir)


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate paper tables and figures")
    parser.add_argument(
        "--results",
        type=Path,
        required=True,
        help="Path to a long-form CSV or Parquet with evaluation results",
    )
    parser.add_argument(
        "--config",
        type=Path,
        default=DEFAULT_CONFIG_PATH,
        help="Path to the paper YAML registry",
    )
    parser.add_argument(
        "--artifacts",
        type=Path,
        default=DEFAULT_ARTIFACTS_DIR,
        help="Output directory for generated artifacts",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    generate_artifacts(args.results, config_path=args.config, artifacts_dir=args.artifacts)


if __name__ == "__main__":
    main()
